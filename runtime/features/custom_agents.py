from __future__ import annotations

import json
import re
from html import escape
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from .base import HomeHubFeature, RuntimeBridge

try:
    from cortex.core import AgentCortex
    from cortex.models import default_agent_cortex
except ModuleNotFoundError:
    from runtime.cortex.core import AgentCortex
    from runtime.cortex.models import default_agent_cortex


ZH = {
    "qs": [
        ("name", "\u8fd9\u4e2a\u667a\u80fd\u4f53\u60f3\u53eb\u4ec0\u4e48\u540d\u5b57\uff1f"),
        ("goal", "\u5b83\u957f\u671f\u8981\u8d1f\u8d23\u4ec0\u4e48\u4efb\u52a1\uff1f"),
        ("primaryUser", "\u5b83\u4e3b\u8981\u670d\u52a1\u8c01\uff1f"),
        ("trigger", "\u5b83\u5e94\u8be5\u5728\u4ec0\u4e48\u65f6\u673a\u5f00\u59cb\u5de5\u4f5c\uff1f"),
        ("inputs", "\u5b83\u5de5\u4f5c\u65f6\u9700\u8981\u54ea\u4e9b\u8d44\u6599\u6216\u8f93\u5165\uff1f"),
        ("output", "\u5b83\u5b8c\u6210\u540e\u8981\u8f93\u51fa\u4ec0\u4e48\u7ed3\u679c\uff1f"),
        ("checkPrompt", "\u5b83\u6bcf\u6b21\u4e3b\u52a8\u68c0\u67e5\u65f6\uff0c\u5e94\u8be5\u600e\u4e48\u95ee\u7528\u6237\uff1f"),
        ("noInputAction", "\u5982\u679c\u7528\u6237\u8bf4\u6ca1\u6709\u65b0\u5185\u5bb9\uff0c\u5b83\u8fd9\u4e00\u8f6e\u5e94\u8be5\u600e\u4e48\u5904\u7406\uff1f"),
        ("hasInputAction", "\u5982\u679c\u7528\u6237\u53d1\u6765\u4e86\u6587\u5b57\u6216\u56fe\u7247\uff0c\u5b83\u8981\u600e\u4e48\u8bb0\u5f55\u548c\u7ee7\u7eed\u5904\u7406\uff1f"),
        ("constraints", "\u5b83\u5fc5\u987b\u9075\u5b88\u4ec0\u4e48\u9650\u5236\u6216\u539f\u5219\uff1f"),
    ],
    "type_name": "\u901a\u7528\u5bb6\u5ead\u667a\u80fd\u4f53",
    "type_summary": "\u6839\u636e\u81ea\u7136\u8bed\u8a00\u9700\u6c42\u521b\u5efa HomeHub \u667a\u80fd\u4f53\uff0c\u5e76\u5728\u7f3a\u8d44\u6599\u65f6\u4e3b\u52a8\u8ffd\u95ee\u3002",
    "type_example": "\u5e2e\u6211\u521b\u5efa\u4e00\u4e2a\u6bcf\u5468\u6574\u7406\u5bb6\u5ead\u8d26\u5355\u7684\u667a\u80fd\u4f53",
    "new_agent_name": "\u65b0\u7684\u5bb6\u5ead\u667a\u80fd\u4f53",
    "default_goal": "\u5e2e\u52a9\u5bb6\u5ead\u5904\u7406\u4e00\u9879\u91cd\u590d\u6027\u4efb\u52a1",
    "scheduled_trigger": "\u5b9a\u65f6\u89e6\u53d1",
    "event_trigger": "\u6536\u5230\u65b0\u8f93\u5165\u540e\u89e6\u53d1",
    "empty": "\u8fd8\u6ca1\u6709\u901a\u7528\u5bb6\u5ead\u667a\u80fd\u4f53\u3002",
    "help": "\u4f60\u53ef\u4ee5\u76f4\u63a5\u63cf\u8ff0\u8fd9\u4e2a\u667a\u80fd\u4f53\u8981\u957f\u671f\u8d1f\u8d23\u4ec0\u4e48\uff0c\u6211\u4f1a\u5728\u7f3a\u8d44\u6599\u65f6\u7ee7\u7eed\u8ffd\u95ee\u4f60\u3002",
    "draft_started": "{name} \u5df2\u5f00\u59cb\u8bbe\u8ba1\u3002{question}",
    "done": "{name} \u7684\u84dd\u56fe\u5df2\u7ecf\u5b8c\u6210\u3002\u5982\u679c\u4f60\u613f\u610f\uff0c\u53ef\u4ee5\u7ee7\u7eed\u8bf4\u201c\u628a{name}\u751f\u6210 feature \u6a21\u677f\u201d\u3002",
    "cancel": "\u597d\u7684\uff0c\u6211\u5148\u505c\u6b62\u5b8c\u5584 {name}\u3002",
    "generated": "{name} \u7684 feature \u6a21\u677f\u5df2\u751f\u6210\uff0c\u8def\u5f84\u662f {path}\u3002",
    "missing_feature": "\u6211\u8fd8\u6ca1\u627e\u5230\u5df2\u5b8c\u6210\u7684\u84dd\u56fe\u53ef\u4ee5\u751f\u6210 feature\u3002",
    "not_ready": "{name} \u8fd8\u5728\u6536\u96c6\u8d44\u6599\uff0c\u8fd8\u4e0d\u80fd\u751f\u6210 feature \u6a21\u677f\u3002",
    "exists": "{name} \u7684 feature \u6587\u4ef6\u5df2\u5b58\u5728\uff1a{path}\u3002\u5982\u679c\u8981\u8986\u76d6\uff0c\u53ef\u4ee5\u518d\u8bf4\u4e00\u6b21\u201c\u8986\u76d6\u751f\u6210 feature\u201d\u3002",
    "detail": "{name} \u5df2\u5b8c\u6210\u3002\u6838\u5fc3\u804c\u8d23\uff1a{goal}\uff1b\u89e6\u53d1\uff1a{trigger}\uff1b\u8f93\u51fa\uff1a{output}\u3002",
    "collecting": "{name} \u8fd8\u5728\u6536\u96c6\u4fe1\u606f\u3002\u4e0b\u4e00\u6b65\u95ee\u9898\u662f\uff1a{question}",
    "review_intro": "{name} \u7684\u521b\u5efa\u84dd\u56fe\u5df2\u7ecf\u6574\u7406\u597d\u4e86\uff0c\u8bf7\u5148\u786e\u8ba4\u8fd9\u4e2a\u603b\u7ed3\uff1a",
    "review_confirm": "\u5982\u679c\u6ca1\u95ee\u9898\uff0c\u4f60\u53ef\u4ee5\u56de\u590d\u201c\u786e\u8ba4\u521b\u5efa\u201d\u3001\u201c\u53ef\u4ee5\u5f00\u59cb\u201d\u6216\u201cOK\u201d\u3002\u5982\u679c\u8981\u4fee\u6539\uff0c\u76f4\u63a5\u544a\u8bc9\u6211\u8981\u6539\u54ea\u91cc\u3002",
    "confirmed": "{name} \u5df2\u6b63\u5f0f\u521b\u5efa\u5b8c\u6210\u3002\u540e\u9762\u6211\u4f1a\u6309\u7167\u8fd9\u4e2a\u84dd\u56fe\u53bb\u63d0\u9192\u3001\u8bb0\u5f55\u548c\u5904\u7406\u65b0\u8f93\u5165\u3002",
    "updated_review": "\u597d\u7684\uff0c\u6211\u5df2\u7ecf\u6839\u636e\u4f60\u7684\u8865\u5145\u66f4\u65b0\u4e86\u84dd\u56fe\uff0c\u8bf7\u518d\u786e\u8ba4\u4e00\u6b21\uff1a",
    "edit_started": "\u6211\u5df2\u7ecf\u57fa\u4e8e {name} \u5f00\u4e86\u4e00\u4e2a\u4fee\u6539\u7248\u8349\u7a3f\u3002\u4f60\u53ef\u4ee5\u7ee7\u7eed\u8bf4\u8981\u4fee\u6539\u54ea\u4e9b\u5185\u5bb9\uff0c\u6bd4\u5982\u540d\u5b57\u3001\u76ee\u6807\u3001\u8f93\u5165\u3001\u8f93\u51fa\u6216\u7ea6\u675f\u3002",
    "edit_confirmed": "{name} \u7684\u4fee\u6539\u7248\u5df2\u786e\u8ba4\u5b8c\u6210\uff0c\u6211\u5df2\u7ecf\u6309\u7167\u65b0\u7684\u5185\u5bb9\u751f\u6210\u4e86\u4e00\u4e2a\u65b0\u667a\u80fd\u4f53\u7248\u672c\u3002",
    "module_summary_count": "\u5df2\u521b\u5efa {total} \u4e2a\u901a\u7528\u667a\u80fd\u4f53\uff0c\u5176\u4e2d {completed} \u4e2a\u5df2\u5b8c\u6210\u84dd\u56fe\u3002",
    "module_summary_empty": "\u8fd8\u6ca1\u6709\u901a\u7528\u667a\u80fd\u4f53\uff0c\u53ef\u4ee5\u5148\u7528\u8bed\u97f3\u6216\u6587\u5b57\u63cf\u8ff0\u4f60\u60f3\u521b\u5efa\u7684\u957f\u671f\u52a9\u624b\u3002",
    "module_generated_suffix": " \u5176\u4e2d {count} \u4e2a\u5df2\u751f\u6210 feature \u811a\u624b\u67b6\u3002",
    "module_collecting_suffix": " \u76ee\u524d\u8fd8\u6709 {count} \u4e2a\u84dd\u56fe\u5728\u8865\u5168\u8d44\u6599\u3002",
    "module_name": "\u667a\u80fd\u4f53\u5de5\u4f5c\u5ba4",
    "module_action_continue": "\u7ee7\u7eed\u5b8c\u5584",
    "module_action_create": "\u5f00\u59cb\u521b\u5efa",
    "artifact_notes": [
        "\u5728\u8bed\u97f3\u6216\u6587\u5b57\u5bf9\u8bdd\u4e2d\u8bc6\u522b\u7528\u6237\u9700\u6c42\u3002",
        "\u5982\u679c\u8d44\u6599\u4e0d\u5b8c\u6574\uff0c\u5148\u4e3b\u52a8\u8ffd\u95ee\uff0c\u518d\u6267\u884c\u3002",
        "\u5728\u5141\u8bb8\u8303\u56f4\u5185\u590d\u7528\u672c\u5730\u5bb6\u5ead\u8bb0\u5fc6\u3001\u6587\u6863\u3001\u65e5\u7a0b\u548c\u6d88\u606f\u6865\u63a5\u8f93\u5165\u3002",
        "\u5728\u6267\u884c\u8fc7\u7a0b\u4e2d\u6301\u7eed\u7ed9\u7528\u6237\u7b80\u77ed\u8fdb\u5ea6\u53cd\u9988\u3002",
    ],
}

AGENT_ACTIVATION_PHRASES = {
    "zh-CN": [
        "\u6253\u5f00\u667a\u80fd\u4f53\u5de5\u4f5c\u5ba4",
        "\u8fdb\u5165\u667a\u80fd\u4f53\u5de5\u4f5c\u5ba4",
        "\u6211\u60f3\u521b\u5efa\u667a\u80fd\u4f53",
        "\u5e2e\u6211\u521b\u5efa\u667a\u80fd\u4f53",
        "\u5e2e\u6211\u505a\u4e2a\u52a9\u624b",
        "\u5217\u51fa\u6211\u7684\u667a\u80fd\u4f53",
        "\u67e5\u770b\u6211\u7684\u667a\u80fd\u4f53",
    ],
    "default": [
        "open custom agents",
        "open agent studio",
        "create an agent",
        "build an assistant",
        "show my agents",
        "list my agents",
    ],
}


def zh(locale: str, text: str, fallback: str) -> str:
    return text if locale == "zh-CN" else fallback


class Feature(HomeHubFeature):
    feature_id = "custom-agents"
    feature_name = "Custom Agents Studio"
    version = "1.0.0"

    def descriptor(self) -> dict:
        data = super().descriptor()
        data["summary"] = "Creates reusable custom household agents, asks follow-up questions, and generates feature scaffolds."
        data["api"] = ["/api/custom-agents", "/api/custom-agents/cortex", "/api/custom-agents/autonomous-plan", "/api/custom-agents/generate-feature", "/api/custom-agents/intake", "/api/custom-agents/lookup"]
        data["voiceIntents"] = self.voice_intents()
        return data

    def voice_intents(self) -> list[dict]:
        return [{
            "id": "custom-agent-builder",
            "name": "Custom Agent Builder",
            "summary": "Creates custom HomeHub agents, routes messages to matching existing agents, and can generate feature scaffolds from completed blueprints.",
            "activationPhrases": AGENT_ACTIVATION_PHRASES["zh-CN"] + AGENT_ACTIVATION_PHRASES["default"],
        }]

    def list_agent_types(self, locale: str, runtime: RuntimeBridge) -> list[dict]:
        if locale == "zh-CN":
            return [{"id": "custom-agent", "name": ZH["type_name"], "summary": ZH["type_summary"], "examplePrompt": ZH["type_example"]}]
        return [{"id": "custom-agent", "name": "Custom Household Agent", "summary": "Turns natural-language requests into reusable HomeHub agents and asks follow-up questions when requirements are incomplete.", "examplePrompt": "Create an agent that reviews our family bills every week"}]

    def storage_path(self, runtime: RuntimeBridge) -> Path:
        agents_dir = runtime.root / "agents"
        agents_dir.mkdir(parents=True, exist_ok=True)
        return agents_dir / "custom_agents.json"

    def cortex(self, runtime: RuntimeBridge) -> AgentCortex:
        return AgentCortex(runtime.root)

    def default_store(self) -> dict:
        now = datetime.now().replace(second=0, microsecond=0).isoformat(timespec="minutes")
        return {"items": [], "recentActions": [{"id": "custom-agents-ready", "summary": "Custom agent studio is ready.", "createdAt": now}]}

    def load_store(self, runtime: RuntimeBridge) -> dict:
        path = self.storage_path(runtime)
        if not path.exists():
            store = self.default_store()
            self.save_store(store, runtime)
            return store
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            store = self.default_store()
            self.save_store(store, runtime)
            return store
        if not isinstance(data, dict):
            return self.default_store()
        return {"items": data.get("items", []) if isinstance(data.get("items", []), list) else [], "recentActions": data.get("recentActions", []) if isinstance(data.get("recentActions", []), list) else []}

    def save_store(self, store: dict, runtime: RuntimeBridge) -> None:
        self.storage_path(runtime).write_text(json.dumps(store, ensure_ascii=False, indent=2), encoding="utf-8")

    def on_refresh(self, runtime: RuntimeBridge) -> None:
        runtime.state[self.feature_id] = self.load_store(runtime)
        cortex = self.cortex(runtime)
        for agent in self.sorted_agents(runtime):
            cortex.sync_agent(agent, stage=str(agent.get("status", "")).strip() or "loaded")

    def reset(self, runtime: RuntimeBridge) -> None:
        store = self.default_store()
        runtime.state[self.feature_id] = store
        self.save_store(store, runtime)

    def get_store(self, runtime: RuntimeBridge) -> dict:
        store = runtime.state.get(self.feature_id)
        if not isinstance(store, dict):
            store = self.load_store(runtime)
            runtime.state[self.feature_id] = store
        return store

    def sorted_agents(self, runtime: RuntimeBridge) -> list[dict]:
        return sorted(self.get_store(runtime).get("items", []), key=lambda item: item.get("updatedAt", ""), reverse=True)

    def enrich_agent_runtime_stats(self, agent: dict, runtime: RuntimeBridge) -> dict:
        enriched = deepcopy(agent)
        generated_feature_id = str(agent.get("generatedFeatureId", "")).strip()
        feature_store = runtime.state.get(generated_feature_id) if generated_feature_id else None
        if isinstance(feature_store, dict):
            items = feature_store.get("items", [])
            if isinstance(items, list):
                enriched["featureItemCount"] = len(items)
            recent = feature_store.get("recentActions", [])
            if isinstance(recent, list) and recent:
                enriched["featureLatestAction"] = str(recent[0].get("summary", ""))
        return enriched

    def append_action(self, runtime: RuntimeBridge, summary: str) -> None:
        store = self.get_store(runtime)
        stamp = datetime.now().replace(second=0, microsecond=0).isoformat(timespec="minutes")
        store.setdefault("recentActions", []).insert(0, {"id": f"custom-act-{stamp}", "summary": summary, "createdAt": stamp})
        del store["recentActions"][12:]

    def is_create_request(self, message: str) -> bool:
        lowered = message.lower()
        cn = any(token in message for token in ["\u521b\u5efa", "\u65b0\u5efa", "\u5e2e\u6211\u505a", "\u751f\u6210"]) and any(token in message for token in ["\u667a\u80fd\u4f53", "\u52a9\u624b", "\u4ee3\u7406", "\u6d41\u7a0b"])
        en = any(token in lowered for token in ["create", "build", "make", "generate"]) and any(token in lowered for token in ["agent", "assistant", "bot", "workflow"])
        return cn or en

    def matches_activation_phrase(self, message: str, locale: str) -> bool:
        normalized = str(message or "").strip().lower()
        if not normalized:
            return False
        phrases = AGENT_ACTIVATION_PHRASES.get(locale, []) + AGENT_ACTIVATION_PHRASES["default"]
        return any(phrase.lower() in normalized for phrase in phrases)

    def is_list_request(self, message: str) -> bool:
        lowered = message.lower()
        return any(token in message for token in ["\u67e5\u770b", "\u5217\u51fa", "\u6709\u54ea\u4e9b", "\u8be6\u60c5"]) or any(token in lowered for token in ["list", "show", "view", "details"])

    def is_cancel_request(self, message: str) -> bool:
        lowered = message.lower()
        return any(token in message for token in ["\u53d6\u6d88", "\u505c\u6b62", "\u4e0d\u7528\u4e86"]) or any(token in lowered for token in ["cancel", "stop", "never mind"])

    def should_force_create(self, message: str) -> bool:
        lowered = message.lower()
        return any(token in message for token in ["\u4ecd\u7136\u65b0\u5efa", "\u8fd8\u662f\u65b0\u5efa", "\u91cd\u65b0\u521b\u5efa", "\u53e6\u5efa\u4e00\u4e2a"]) or any(
            token in lowered for token in ["create anyway", "new one anyway", "create another", "still create"]
        )

    def looks_like_agent_topic(self, message: str) -> bool:
        lowered = message.lower()
        return any(token in message for token in ["\u667a\u80fd\u4f53", "\u52a9\u624b", "\u4ee3\u7406", "\u673a\u5668\u4eba", "\u6d41\u7a0b"]) or any(token in lowered for token in ["agent", "assistant", "bot", "workflow"])

    def should_generate_feature(self, message: str) -> bool:
        lowered = message.lower()
        return "feature" in lowered and (any(token in message for token in ["\u751f\u6210", "\u8986\u76d6"]) or any(token in lowered for token in ["generate", "template", "scaffold", "overwrite"]))

    def should_overwrite_feature(self, message: str) -> bool:
        lowered = message.lower()
        return "\u8986\u76d6" in message or "overwrite" in lowered or "replace" in lowered

    def get_collecting_agent(self, runtime: RuntimeBridge) -> dict | None:
        for agent in self.sorted_agents(runtime):
            if agent.get("status") == "collecting":
                return agent
        return None

    def get_review_agent(self, runtime: RuntimeBridge) -> dict | None:
        for agent in self.sorted_agents(runtime):
            if agent.get("status") == "review":
                return agent
        return None

    def latest_completed_agent(self, runtime: RuntimeBridge) -> dict | None:
        for agent in self.sorted_agents(runtime):
            if agent.get("status") == "complete":
                return agent
        return None

    def latest_operational_agent(self, runtime: RuntimeBridge) -> dict | None:
        for agent in self.sorted_agents(runtime):
            if agent.get("status") == "complete" and any(str(agent.get("profile", {}).get(key, "")).strip() for key in ["checkPrompt", "noInputAction", "hasInputAction"]):
                return agent
        return None

    def find_agent_by_hint(self, message: str, runtime: RuntimeBridge) -> dict | None:
        lowered = message.lower()
        for agent in self.sorted_agents(runtime):
            if agent.get("name") and agent["name"] in message:
                return agent
            goal = str(agent.get("profile", {}).get("goal", "")).lower()
            if goal and goal[:24] in lowered:
                return agent
        return None

    def split_keywords(self, text: str) -> set[str]:
        raw = str(text or "").lower()
        return {
            token.strip()
            for token in re.split(r"[\s,，。；;、:/|()\[\]{}<>\"'\n\r\t\-]+", raw)
            if len(token.strip()) >= 2
        }

    def keyword_set(self, text: str) -> set[str]:
        raw = str(text or "").lower()
        return {
            token.strip()
            for token in re.split(r"[\s,，。；;、:/|()\[\]{}<>\"'\n\r\t]+", raw)
            if len(token.strip()) >= 2
        }

    def split_keywords(self, text: str) -> set[str]:
        raw = str(text or "").lower()
        tokens = {
            token.strip()
            for token in re.split(r"[\s,/|()\[\]{}<>\"'\n\r\t\-]+", raw)
            if len(token.strip()) >= 2
        }
        for chunk in re.findall(r"[\u4e00-\u9fff]{2,}", raw):
            size = len(chunk)
            for width in range(2, min(5, size + 1)):
                for start in range(0, size - width + 1):
                    tokens.add(chunk[start:start + width])
        return tokens

    def agent_business_keywords(self, agent: dict) -> set[str]:
        profile = agent.get("profile", {})
        autonomous = agent.get("autonomousCreation", {}) if isinstance(agent.get("autonomousCreation", {}), dict) else {}
        parts = [
            agent.get("name", ""),
            profile.get("goal", ""),
            profile.get("inputs", ""),
            profile.get("output", ""),
            profile.get("checkPrompt", ""),
            profile.get("hasInputAction", ""),
            profile.get("constraints", ""),
            autonomous.get("createdFromDemand", ""),
        ]
        keywords: set[str] = set()
        for part in parts:
            keywords.update(self.split_keywords(str(part)))
        return keywords

    def business_match_score(self, agent: dict, message: str) -> float:
        message_tokens = self.split_keywords(message)
        message_text = str(message or "").strip().lower()
        if not message_tokens:
            message_tokens = set()
        keywords = self.agent_business_keywords(agent)
        profile = agent.get("profile", {})
        fields = [
            str(agent.get("name", "")).strip().lower(),
            str(profile.get("goal", "")).strip().lower(),
            str(profile.get("inputs", "")).strip().lower(),
            str(profile.get("output", "")).strip().lower(),
            str(profile.get("checkPrompt", "")).strip().lower(),
            str(profile.get("hasInputAction", "")).strip().lower(),
            str((agent.get("autonomousCreation", {}) if isinstance(agent.get("autonomousCreation", {}), dict) else {}).get("createdFromDemand", "")).strip().lower(),
        ]
        if not keywords and not any(fields):
            return 0.0
        overlap = message_tokens & keywords
        score = len(overlap)
        if agent.get("name") and str(agent.get("name")) in message:
            score += 4
        for field in fields:
            if not field:
                continue
            if field in message_text:
                score += 3
                continue
            for token in self.split_keywords(field):
                if len(token) >= 2 and token in message_text:
                    score += 1
        goal = str(profile.get("goal", "")).strip()
        if goal and any(token in goal.lower() for token in message_tokens):
            score += 1
        return float(score)

    def find_matching_operational_agent(self, message: str, runtime: RuntimeBridge) -> dict | None:
        candidates: list[tuple[float, dict]] = []
        for agent in self.sorted_agents(runtime):
            if agent.get("status") != "complete":
                continue
            if not any(str(agent.get("profile", {}).get(key, "")).strip() for key in ["checkPrompt", "noInputAction", "hasInputAction"]):
                continue
            score = self.business_match_score(agent, message)
            if score > 0:
                candidates.append((score, agent))
        if not candidates:
            return None
        candidates.sort(key=lambda item: (item[0], item[1].get("updatedAt", "")), reverse=True)
        return candidates[0][1] if candidates[0][0] >= 1 else None

    def similarity_score(self, left: dict, right: dict) -> float:
        left_tokens = self.agent_business_keywords(left)
        right_tokens = self.agent_business_keywords(right)
        if not left_tokens or not right_tokens:
            return 0.0
        overlap = left_tokens & right_tokens
        union = left_tokens | right_tokens
        score = len(overlap) / max(len(union), 1)
        if left.get("name") and right.get("name") and left.get("name") == right.get("name"):
            score += 0.35
        return score

    def find_similar_agents(self, draft_profile: dict, runtime: RuntimeBridge, limit: int = 3) -> list[dict]:
        draft = {"name": draft_profile.get("name", ""), "profile": draft_profile}
        ranked: list[tuple[float, dict]] = []
        for agent in self.sorted_agents(runtime):
            score = self.similarity_score(draft, agent)
            if score >= 0.18:
                ranked.append((score, agent))
        ranked.sort(key=lambda item: (item[0], item[1].get("updatedAt", "")), reverse=True)
        return [item[1] for item in ranked[:limit]]

    def similar_agents_message(self, draft_profile: dict, matches: list[dict], locale: str) -> str:
        if not matches:
            return ""
        names = "、".join(str(item.get("name", "")) for item in matches if item.get("name"))
        if locale == "zh-CN":
            requested = draft_profile.get("name") or draft_profile.get("goal") or "\u8fd9\u4e2a\u667a\u80fd\u4f53"
            return f"\u6211\u627e\u5230\u4e86\u51e0\u4e2a\u76f8\u8fd1\u7684\u667a\u80fd\u4f53\uff1a{names}\u3002\u5982\u679c\u4f60\u662f\u60f3\u7ee7\u7eed\u7528\u5b83\u4eec\uff0c\u53ef\u4ee5\u76f4\u63a5\u8bf4\u540d\u5b57\uff1b\u5982\u679c\u4ecd\u7136\u8981\u65b0\u5efa {requested}\uff0c\u8bf7\u660e\u786e\u8bf4\u201c\u4ecd\u7136\u65b0\u5efa\u201d\u3002"
        return f"I found similar agents: {names}. If you want one of those, say its name. If you still want a brand new one, say 'create anyway'."

    def next_question(self, agent: dict, locale: str) -> tuple[str, str] | None:
        profile = agent.get("profile", {})
        for key, prompt in ZH["qs"]:
            value = str(profile.get(key, "")).strip()
            if key == "trigger" and value in {ZH["scheduled_trigger"], ZH["event_trigger"], "scheduled trigger", "event-driven trigger"}:
                return key, zh(locale, prompt, f"Please provide {key}.")
            if not value:
                return key, zh(locale, prompt, f"Please provide {key}.")
        return None

    def is_confirmation_message(self, message: str) -> bool:
        lowered = message.lower().strip()
        confirm_tokens = [
            "\u786e\u8ba4\u521b\u5efa",
            "\u786e\u8ba4\u4fee\u6539",
            "\u53ef\u4ee5\u5f00\u59cb",
            "\u5f00\u59cb\u521b\u5efa",
            "\u6ca1\u95ee\u9898",
            "\u597d\u7684",
            "\u786e\u8ba4",
            "\u5b8c\u6210\u4fee\u6539",
            "\u4fee\u6539\u5b8c\u6210",
            "\u5c31\u8fd9\u6837",
        ]
        return any(token in message for token in confirm_tokens) or lowered in {"ok", "okay", "confirm", "looks good", "go ahead"}

    def is_revision_message(self, message: str) -> bool:
        lowered = message.lower()
        revision_tokens = [
            "\u4fee\u6539",
            "\u6539\u4e00\u4e0b",
            "\u8c03\u6574",
            "\u8865\u5145",
            "\u4fee\u6b63",
            "\u6539\u6210",
        ]
        return any(token in message for token in revision_tokens) or any(token in lowered for token in ["change", "update", "adjust", "revise"])

    def extract_partial_blueprint(self, message: str, locale: str, runtime: RuntimeBridge) -> dict:
        payload = runtime.openai_json(
            "You extract HomeHub custom-agent requirements from a natural-language request. Return JSON only with keys: shouldCreateAgent, name, goal, primaryUser, trigger, inputs, output, constraints, checkPrompt, noInputAction, hasInputAction, allowNetworkLookup, preferredSources, lookupPolicy, confidence. Use empty strings for unknown values.",
            json.dumps({"locale": locale, "message": message}, ensure_ascii=False),
            "gpt-4o-mini",
        )
        return payload if isinstance(payload, dict) else {}

    def classify_agent_message(self, message: str, locale: str, runtime: RuntimeBridge) -> dict:
        agents = [
            {
                "name": str(item.get("name", "")).strip(),
                "status": str(item.get("status", "")).strip(),
                "goal": str(item.get("profile", {}).get("goal", "")).strip(),
                "trigger": str(item.get("profile", {}).get("trigger", "")).strip(),
                "generatedFeatureId": str(item.get("generatedFeatureId", "")).strip(),
            }
            for item in self.sorted_agents(runtime)[:8]
        ]
        context = {
            "locale": locale,
            "message": message,
            "collectingAgent": (self.get_collecting_agent(runtime) or {}).get("name", ""),
            "reviewAgent": (self.get_review_agent(runtime) or {}).get("name", ""),
            "latestOperationalAgent": (self.latest_operational_agent(runtime) or {}).get("name", ""),
            "agents": agents,
        }
        payload = runtime.openai_json(
            (
                "You are the HomeHub custom-agent intent classifier. "
                "Decide whether the user's message is about creating a custom agent, continuing blueprint collection, "
                "reviewing a blueprint, editing an existing custom agent, operating an existing custom agent, showing an agent, or just asking for general help. "
                "Return JSON only with keys: action, targetAgentName, confidence, reason. "
                "Valid action values: create_agent, continue_collecting, review_blueprint, edit_agent, operational_agent, show_agent, general_agent_help, no_match. "
                "Use no_match when this should be handled by another HomeHub system such as schedule, reminders, weather, or general chat."
            ),
            json.dumps(context, ensure_ascii=False),
            "gpt-4o-mini",
        )
        return payload if isinstance(payload, dict) else {}

    def default_name(self, goal: str, locale: str) -> str:
        trimmed = re.sub(r"\s+", " ", goal).strip(" .;,\u3002\uff0c\uff1b")
        if not trimmed:
            return ZH["new_agent_name"] if locale == "zh-CN" else "New HomeHub Agent"
        return f"{trimmed[:18]}\u667a\u80fd\u4f53" if locale == "zh-CN" else f"{trimmed[:28]} Agent"

    def infer_initial_profile(self, message: str, locale: str, runtime: RuntimeBridge) -> dict:
        ai = self.extract_partial_blueprint(message, locale, runtime)
        lowered = message.lower()
        keys = ["name", "goal", "primaryUser", "trigger", "inputs", "output", "constraints", "checkPrompt", "noInputAction", "hasInputAction", "allowNetworkLookup", "preferredSources", "lookupPolicy"]
        profile = {key: str(ai.get(key, "")).strip() for key in keys}
        if not profile["goal"]:
            profile["goal"] = message.strip() or zh(locale, ZH["default_goal"], "Help with an ongoing household task")
        if not profile["trigger"]:
            if any(token in message for token in ["\u6bcf\u5929", "\u6bcf\u5468", "\u6bcf\u6708"]):
                profile["trigger"] = zh(locale, ZH["scheduled_trigger"], "scheduled trigger")
            elif any(token in lowered for token in ["daily", "weekly", "monthly"]):
                profile["trigger"] = "scheduled trigger"
            elif any(token in message for token in ["\u6536\u5230", "\u4e0a\u4f20", "\u65b0\u6d88\u606f"]):
                profile["trigger"] = zh(locale, ZH["event_trigger"], "event-driven trigger")
        if not profile["name"]:
            match = re.search(r"\u53eb(?:\u505a)?(.{2,20}?)(?:\u7684)?(?:\u667a\u80fd\u4f53|\u52a9\u624b|\u4ee3\u7406|\u673a\u5668\u4eba)", message)
            if match:
                profile["name"] = match.group(1).strip()
        if any(token in message for token in ["\u8d26\u5355", "\u6263\u8d39", "\u53d1\u7968", "\u6536\u636e", "\u622a\u56fe"]) or any(token in lowered for token in ["bill", "invoice", "expense", "charge", "receipt", "screenshot"]):
            profile["inputs"] = profile["inputs"] or zh(locale, "\u8d26\u5355\u622a\u56fe\u3001\u8d26\u5355\u6587\u5b57\u3001\u6263\u8d39\u6d88\u606f", "Bill screenshots, bill text, and charge notifications")
            profile["output"] = profile["output"] or zh(locale, "\u8d26\u5355\u6c47\u603b\u3001\u5f02\u5e38\u63d0\u9192\u3001\u8bb0\u5f55\u5f52\u6863", "Bill summaries, anomaly alerts, and archived records")
            profile["checkPrompt"] = profile["checkPrompt"] or zh(locale, "\u4eca\u5929\u6709\u65b0\u7684\u8d26\u5355\u6216\u6263\u8d39\u8bb0\u5f55\u5417\uff1f", "Do you have any new bills or charge records today?")
            profile["noInputAction"] = profile["noInputAction"] or zh(locale, "\u8bb0\u5f55\u672c\u6b21\u65e0\u65b0\u8d26\u5355\uff0c\u7ed3\u675f\u8fd9\u4e00\u8f6e\u68c0\u67e5", "Record that there are no new bills this round and finish the check")
            profile["hasInputAction"] = profile["hasInputAction"] or zh(locale, "\u63a5\u6536\u8d26\u5355\u56fe\u7247\u6216\u6587\u5b57\uff0c\u8bb0\u5f55\u53d1\u9001\u65f6\u95f4\uff0c\u5f52\u6863\u5185\u5bb9\u5e76\u8f93\u51fa\u6458\u8981", "Accept bill images or text, record the send time, archive the content, and output a summary")
            profile["constraints"] = profile["constraints"] or zh(locale, "\u5728\u786e\u8ba4\u524d\u4e0d\u8981\u81ea\u52a8\u8ba4\u5b9a\u6263\u6b3e\u6b63\u5e38", "Do not assume charges are valid before confirmation")
        profile["allowNetworkLookup"] = profile["allowNetworkLookup"] or "yes"
        if any(token in message for token in ["\u67e5\u4e00\u4e0b", "\u67e5\u8be2", "\u641c\u7d22", "\u8054\u7f51", "\u5b98\u65b9", "\u6700\u65b0", "\u65b0\u95fb", "\u5929\u6c14"]) or any(
            token in lowered for token in ["search", "lookup", "web", "online", "official", "latest", "news", "weather"]
        ):
            profile["preferredSources"] = profile["preferredSources"] or "official websites, docs, standards, github.com, huggingface.co"
        else:
            profile["preferredSources"] = profile["preferredSources"] or "official websites, docs, github.com, huggingface.co"
        profile["lookupPolicy"] = profile["lookupPolicy"] or "official-only"
        return profile

    def extract_update_patch(self, message: str, locale: str, runtime: RuntimeBridge) -> dict:
        ai = self.extract_partial_blueprint(message, locale, runtime)
        patch = {key: str(ai.get(key, "")).strip() for key in ["name", "goal", "primaryUser", "trigger", "inputs", "output", "constraints", "checkPrompt", "noInputAction", "hasInputAction", "allowNetworkLookup", "preferredSources", "lookupPolicy"] if str(ai.get(key, "")).strip()}
        lowered = message.lower()
        if not patch.get("trigger") and (
            any(token in message for token in ["\u6bcf\u5929", "\u6bcf\u5468", "\u6bcf\u6708", "\u5468\u4e00", "\u5468\u65e5"])
            or any(token in lowered for token in ["daily", "weekly", "monthly", "monday", "tuesday", "sunday"])
        ):
            patch["trigger"] = message.strip()
        if not patch.get("constraints") and any(token in message for token in ["\u4e0d\u8981", "\u5fc5\u987b", "\u7981\u6b62", "\u9650\u5236", "\u539f\u5219"]):
            patch["constraints"] = message.strip()
        if not patch.get("output") and any(token in message for token in ["\u8f93\u51fa", "\u7ed3\u679c", "\u603b\u7ed3", "\u63d0\u9192"]):
            patch["output"] = message.strip()
        if any(token in message for token in ["\u67e5\u4e00\u4e0b", "\u67e5\u8be2", "\u641c\u7d22", "\u8054\u7f51", "\u5b98\u65b9", "\u6700\u65b0"]) or any(token in lowered for token in ["search", "lookup", "web", "online", "official"]):
            patch["allowNetworkLookup"] = "yes"
            patch.setdefault("preferredSources", "official websites, docs, github.com, huggingface.co")
            patch.setdefault("lookupPolicy", "official-only")
        patch.setdefault("allowNetworkLookup", "yes")
        patch.setdefault("lookupPolicy", "official-only")
        return patch

    def blueprint_allows_network(self, agent: dict) -> bool:
        value = str(agent.get("profile", {}).get("allowNetworkLookup", "")).strip().lower()
        return not value or value in {"yes", "true", "allowed", "allow", "on", "1"}

    def preferred_sources(self, agent: dict) -> list[str]:
        raw = str(agent.get("profile", {}).get("preferredSources", "")).strip()
        if not raw:
            return []
        return [item.strip() for item in re.split(r"[,，;\n]+", raw) if item.strip()]

    def is_expense_or_bill_agent(self, agent: dict, message: str) -> bool:
        haystacks = [
            str(agent.get("name", "")),
            str(agent.get("generatedFeatureId", "")),
            str(agent.get("profile", {}).get("goal", "")),
            str(agent.get("profile", {}).get("inputs", "")),
            str(agent.get("profile", {}).get("output", "")),
            message,
        ]
        combined = " ".join(haystacks).lower()
        tokens = ["账单", "消费", "费用", "支出", "记账", "明细", "bill", "expense", "receipt", "receipt", "budget"]
        return any(token in combined for token in tokens)

    def extract_expenses_from_analysis(self, analysis: dict, message: str) -> list[dict]:
        extracted: list[dict] = []
        raw_items = analysis.get("detectedExpenses", [])
        if isinstance(raw_items, list):
            for item in raw_items:
                if not isinstance(item, dict):
                    continue
                try:
                    amount = int(float(item.get("amount", 0) or 0))
                except (TypeError, ValueError):
                    amount = 0
                if amount <= 0:
                    continue
                extracted.append({
                    "amount": amount,
                    "category": str(item.get("category", "")).strip() or "其他",
                    "content": str(item.get("content", "")).strip(),
                    "note": str(item.get("note", "")).strip(),
                    "merchant": str(item.get("merchant", "")).strip() or str(analysis.get("merchant", "")).strip(),
                    "purchaseDate": str(item.get("purchaseDate", "")).strip() or str(analysis.get("detectedDate", "")).strip(),
                    "paymentMethod": str(item.get("paymentMethod", "")).strip() or str(analysis.get("paymentMethod", "")).strip(),
                    "taxAmount": int(float(item.get("taxAmount", analysis.get("taxAmount", 0)) or 0)),
                })
        if extracted:
            return extracted
        facts = analysis.get("relevantFacts", [])
        joined = "\n".join(str(item) for item in facts if str(item).strip())
        matches = re.findall(r"(\d+(?:\.\d+)?)\s*(?:日元|円|yen|jpy)", joined.lower())
        if not matches:
            return []
        default_content = message.strip() or str(analysis.get("summary", "")).strip() or "Image-scanned expense"
        for value in matches:
            try:
                amount = int(float(value))
            except (TypeError, ValueError):
                continue
            if amount <= 0:
                continue
            extracted.append({
                "amount": amount,
                "category": "其他",
                "content": default_content,
                "note": "",
                "merchant": str(analysis.get("merchant", "")).strip(),
                "purchaseDate": str(analysis.get("detectedDate", "")).strip(),
                "paymentMethod": str(analysis.get("paymentMethod", "")).strip(),
                "taxAmount": int(float(analysis.get("taxAmount", 0) or 0)),
            })
        return extracted

    def create_agent(self, runtime: RuntimeBridge, message: str, locale: str) -> dict:
        profile = self.infer_initial_profile(message, locale, runtime)
        name = profile.get("name") or self.default_name(profile.get("goal", ""), locale)
        stamp = datetime.now().strftime("%Y%m%d%H%M%S")
        agent = {"id": f"custom-agent-{stamp}", "type": "custom-agent", "name": name, "status": "collecting", "createdAt": self.now_iso(), "updatedAt": self.now_iso(), "profile": profile, "qaHistory": [], "artifact": "", "generatedFeaturePath": "", "generatedFeatureId": "", "records": []}
        store = self.get_store(runtime)
        store.setdefault("items", []).insert(0, agent)
        self.append_action(runtime, f"Created custom agent draft '{name}'.")
        self.save_store(store, runtime)
        self.cortex(runtime).record_event(agent, "draft_created", {"message": message.strip()})
        return agent

    def build_request_context(
        self,
        message: str,
        locale: str,
        attachments: list[dict] | None = None,
        task_type: str = "general_chat",
        requires_network: bool = False,
        require_artifacts: bool = False,
    ) -> dict:
        attachment_list = attachments if isinstance(attachments, list) else []
        input_modes = ["text"]
        if attachment_list:
            input_modes.append("image")
        lowered = str(message or "").lower()
        if any(token in lowered for token in ["voice", "audio"]) or any(token in str(message or "") for token in ["语音", "音频", "声音"]):
            input_modes.append("voice")
        return {
            "command": str(message or "").strip(),
            "locale": str(locale or "zh-CN").strip() or "zh-CN",
            "taskType": str(task_type or "general_chat").strip() or "general_chat",
            "inputModes": sorted(set(input_modes)),
            "requireArtifacts": bool(require_artifacts or attachment_list),
            "requiresNetwork": bool(requires_network),
            "speakReply": False,
        }

    def best_existing_agent_for_request(self, request: dict, runtime: RuntimeBridge) -> dict | None:
        scored: list[tuple[float, dict, dict]] = []
        message = str(request.get("command", "")).strip()
        for agent in self.sorted_agents(runtime):
            if agent.get("status") != "complete":
                continue
            agent_id = str(agent.get("id", "")).strip()
            if not agent_id:
                continue
            brain = self.cortex(runtime).get_brain(agent_id, request)
            plan = brain.get("autonomousCreation", {}) if isinstance(brain, dict) else {}
            fit = plan.get("existingBrainFit", {}) if isinstance(plan, dict) else {}
            try:
                capability_score = float(fit.get("capabilityFitScore", 0.0) or 0.0)
            except (TypeError, ValueError):
                capability_score = 0.0
            business_score = self.business_match_score(agent, message)
            score = capability_score + min(1.0, business_score / 6.0)
            scored.append((score, agent, brain if isinstance(brain, dict) else {}))
        if not scored:
            return None
        scored.sort(key=lambda item: (item[0], item[1].get("updatedAt", "")), reverse=True)
        best_score, best_agent, best_brain = scored[0]
        return {
            "score": best_score,
            "agent": best_agent,
            "brain": best_brain,
        }

    def should_autonomously_create(self, request: dict, runtime: RuntimeBridge) -> dict:
        command = str(request.get("command", "")).strip()
        if not command:
            return {"shouldCreate": False, "reason": "empty-request"}
        seed_state = default_agent_cortex("homehub-shared-brain", "HomeHub Shared Brain")
        architect = self.cortex(runtime).architect
        design = architect.design_state(seed_state, request)
        autonomous = design.get("autonomousCreation", {}) if isinstance(design, dict) else {}
        requirement = autonomous.get("requirement", {}) if isinstance(autonomous, dict) else {}
        capabilities = requirement.get("requiredCapabilities", []) if isinstance(requirement, dict) else []
        best_existing = self.best_existing_agent_for_request(request, runtime)
        lowered = command.lower()
        continuation_hint = any(token in command for token in ["继续", "沿用", "还是按刚才", "接着处理", "继续用"]) or any(
            token in lowered for token in ["continue", "reuse", "same way", "same workflow", "keep using"]
        )
        if best_existing and (
            float(best_existing.get("score", 0.0) or 0.0) >= 0.75
            or (continuation_hint and float(best_existing.get("score", 0.0) or 0.0) >= 0.35)
        ):
            return {
                "shouldCreate": False,
                "reason": "existing-agent-fit",
                "bestExisting": best_existing,
                "design": design,
            }
        meaningful_capabilities = [item for item in capabilities if item != "semantic-understanding"]
        should_create = bool(meaningful_capabilities) or str(request.get("taskType", "")).strip() in {"agent_creation", "document_analysis", "ocr_analysis"}
        return {
            "shouldCreate": should_create,
            "reason": "capability-gap" if should_create else "general-chat-fit",
            "bestExisting": best_existing,
            "design": design,
        }

    def create_agent_from_autonomous_plan(self, runtime: RuntimeBridge, message: str, locale: str, plan: dict) -> dict:
        design = plan.get("design", {}) if isinstance(plan, dict) else {}
        autonomous = design.get("autonomousCreation", {}) if isinstance(design, dict) else {}
        proposed = autonomous.get("proposedBrain", {}) if isinstance(autonomous, dict) else {}
        profile = self.infer_initial_profile(message, locale, runtime)
        capability_labels = ", ".join(proposed.get("requiredCapabilities", [])) if isinstance(proposed.get("requiredCapabilities", []), list) else ""
        feature_hooks = ", ".join(proposed.get("featureHooks", [])) if isinstance(proposed.get("featureHooks", []), list) else ""
        if not profile.get("name"):
            profile["name"] = self.default_name(str(proposed.get("mission", "")).strip() or profile.get("goal", ""), locale)
        if not profile.get("goal"):
            profile["goal"] = str(proposed.get("mission", "")).strip() or message.strip()
        if not profile.get("inputs"):
            input_modes = proposed.get("inputModes", [])
            if isinstance(input_modes, list) and input_modes:
                profile["inputs"] = ", ".join(str(item) for item in input_modes if str(item).strip())
        if not profile.get("output"):
            profile["output"] = zh(locale, "先完成需求处理，再返回结果、建议或可交付物。", "Complete the request, then return results, guidance, or generated artifacts.")
        profile["allowNetworkLookup"] = str(profile.get("allowNetworkLookup", "")).strip() or "yes"
        profile["preferredSources"] = str(profile.get("preferredSources", "")).strip() or "official websites, docs, github.com, huggingface.co"
        profile["lookupPolicy"] = str(profile.get("lookupPolicy", "")).strip() or "official-only"
        auto_constraints = []
        if capability_labels:
            auto_constraints.append(f"Required capabilities: {capability_labels}")
        if feature_hooks:
            auto_constraints.append(f"Runtime hooks: {feature_hooks}")
        if auto_constraints:
            merged = " | ".join(auto_constraints)
            profile["constraints"] = f"{profile.get('constraints', '').strip()} | {merged}".strip(" |")
        stamp = datetime.now().strftime("%Y%m%d%H%M%S")
        agent = {
            "id": f"custom-agent-{stamp}",
            "type": str(proposed.get("agentType", "custom-agent")).strip() or "custom-agent",
            "name": profile["name"],
            "status": "review",
            "createdAt": self.now_iso(),
            "updatedAt": self.now_iso(),
            "profile": profile,
            "qaHistory": [],
            "artifact": "",
            "generatedFeaturePath": "",
            "generatedFeatureId": "",
            "records": [],
            "autonomousCreation": {
                "createdFromDemand": message.strip(),
                "requirement": autonomous.get("requirement", {}),
                "proposedBrain": proposed,
            },
        }
        agent["artifact"] = self.build_artifact(agent, locale)
        store = self.get_store(runtime)
        store.setdefault("items", []).insert(0, agent)
        self.append_action(runtime, f"Autonomously prepared agent draft '{agent['name']}' from a detected capability gap.")
        self.save_store(store, runtime)
        self.cortex(runtime).record_event(agent, "autonomous_draft_created", {"message": message.strip(), "taskType": str(proposed.get("agentType", "")).strip()})
        return agent

    def maybe_autonomous_agent_response(
        self,
        message: str,
        locale: str,
        runtime: RuntimeBridge,
        task_type: str = "general_chat",
        requires_network: bool = False,
        attachments: list[dict] | None = None,
    ) -> dict | None:
        request = self.build_request_context(
            message,
            locale,
            attachments=attachments,
            task_type=task_type,
            requires_network=requires_network,
            require_artifacts=bool(attachments),
        )
        plan = self.should_autonomously_create(request, runtime)
        best_existing = plan.get("bestExisting") if isinstance(plan, dict) else None
        lowered = str(message or "").lower()
        continuation_hint = any(token in str(message or "") for token in ["继续", "沿用", "还是按刚才", "接着处理", "继续用"]) or any(
            token in lowered for token in ["continue", "reuse", "same way", "same workflow", "keep using"]
        )
        if isinstance(best_existing, dict):
            agent = best_existing.get("agent")
            try:
                score = float(best_existing.get("score", 0.0) or 0.0)
            except (TypeError, ValueError):
                score = 0.0
            if isinstance(agent, dict) and (score >= 0.75 or (continuation_hint and score >= 0.35)):
                reply = zh(
                    locale,
                    f"我判断现有智能体“{agent.get('name', '')}”已经能处理这类需求，先直接复用它来继续执行。",
                    f"I found that the existing agent '{agent.get('name', '')}' already fits this request, so HomeHub will reuse it first.",
                )
                return {
                    "reply": reply + "\n\n" + self.agent_detail_message(agent, locale),
                    "uiAction": self.studio_ui_action(agent),
                    "autonomousPlan": plan,
                    "autonomousAction": "reuse_existing",
                }
        if not plan.get("shouldCreate"):
            return None
        agent = self.create_agent_from_autonomous_plan(runtime, message, locale, plan)
        intro = zh(
            locale,
            f"我先做了能力缺口判断，发现现有智能体还不够匹配，所以我已经为这条需求自主起草了一个新的智能体：{agent['name']}。",
            f"HomeHub detected a capability gap for this request, so it drafted a new agent automatically: {agent['name']}.",
        )
        return {
            "reply": intro + "\n\n" + self.build_review_summary(agent, locale),
            "uiAction": self.studio_ui_action(agent),
            "autonomousPlan": plan,
            "autonomousAction": "create_draft",
        }

    def now_iso(self) -> str:
        return datetime.now().replace(second=0, microsecond=0).isoformat(timespec="minutes")

    def build_artifact(self, agent: dict, locale: str) -> str:
        p = agent.get("profile", {})
        lines = [
            f"# {p.get('name') or agent.get('name')}",
            "",
            "## Goal",
            p.get("goal", ""),
            "",
            "## Primary User",
            p.get("primaryUser", ""),
            "",
            "## Trigger",
            p.get("trigger", ""),
            "",
            "## Required Inputs",
            p.get("inputs", ""),
            "",
            "## Output",
            p.get("output", ""),
            "",
            "## Proactive Check Prompt",
            p.get("checkPrompt", ""),
            "",
            "## If No New Input",
            p.get("noInputAction", ""),
            "",
            "## If New Input Arrives",
            p.get("hasInputAction", ""),
            "",
            "## Network Lookup",
            f"Allowed: {p.get('allowNetworkLookup', '')}",
            f"Preferred Sources: {p.get('preferredSources', '')}",
            f"Lookup Policy: {p.get('lookupPolicy', '')}",
            "",
            "## Constraints",
            p.get("constraints", ""),
            "",
            "## HomeHub Execution Notes",
        ]
        notes = ZH["artifact_notes"] if locale == "zh-CN" else ["Detect the request through voice or text conversation.", "Ask for missing information before execution.", "Reuse local household memory when allowed.", "Keep the user informed with concise progress updates."]
        lines.extend(f"- {item}" for item in notes)
        return "\n".join(lines).strip()

    def complete_agent(self, agent: dict, runtime: RuntimeBridge, locale: str) -> None:
        if not str(agent.get("profile", {}).get("name", "")).strip():
            generated = self.default_name(agent.get("profile", {}).get("goal", ""), locale)
            agent["profile"]["name"] = generated
            agent["name"] = generated
        agent["status"] = "complete"
        agent["updatedAt"] = self.now_iso()
        agent["artifact"] = self.build_artifact(agent, locale)
        self.append_action(runtime, f"Completed custom agent '{agent['name']}'.")
        self.save_store(self.get_store(runtime), runtime)
        self.cortex(runtime).record_event(agent, "confirmed", {"message": agent.get("artifact", "")})

    def is_edit_request(self, message: str) -> bool:
        lowered = str(message or "").lower()
        return (
            any(token in message for token in ["\u4fee\u6539\u667a\u80fd\u4f53", "\u6539\u667a\u80fd\u4f53", "\u4fee\u6539\u8fd9\u4e2a\u667a\u80fd\u4f53", "\u8c03\u6574\u8fd9\u4e2a\u667a\u80fd\u4f53", "\u4fee\u6539\u540d\u5b57", "\u6539\u540d\u5b57"])
            or any(token in lowered for token in ["edit agent", "modify agent", "rename agent", "update this agent"])
        )

    def start_editing_agent(self, base_agent: dict, message: str, runtime: RuntimeBridge, locale: str) -> dict:
        profile = deepcopy(base_agent.get("profile", {}))
        stamp = datetime.now().strftime("%Y%m%d%H%M%S")
        agent = {
            "id": f"custom-agent-{stamp}",
            "type": "custom-agent",
            "name": str(base_agent.get("name") or profile.get("name") or self.default_name(profile.get("goal", ""), locale)).strip(),
            "status": "review",
            "createdAt": self.now_iso(),
            "updatedAt": self.now_iso(),
            "profile": profile,
            "qaHistory": deepcopy(base_agent.get("qaHistory", [])),
            "artifact": "",
            "generatedFeaturePath": "",
            "generatedFeatureId": "",
            "records": [],
            "editMode": True,
            "sourceAgentId": str(base_agent.get("id", "")),
            "sourceAgentName": str(base_agent.get("name", "")),
        }
        patch = self.extract_update_patch(message, locale, runtime)
        for key, value in patch.items():
            if value:
                agent["profile"][key] = value
                if key == "name":
                    agent["name"] = value
        agent["artifact"] = self.build_artifact(agent, locale)
        store = self.get_store(runtime)
        store.setdefault("items", []).insert(0, agent)
        self.append_action(runtime, f"Started edit draft for '{base_agent.get('name', '')}' as '{agent['name']}'.")
        self.save_store(store, runtime)
        self.cortex(runtime).record_event(agent, "edit_started", {"message": message.strip(), "sourceAgentId": str(base_agent.get("id", ""))})
        return agent

    def build_review_summary(self, agent: dict, locale: str) -> str:
        p = agent.get("profile", {})
        mission_label = zh(locale, "长期任务", "Long-term mission")
        user_label = zh(locale, "主要服务对象", "Primary user")
        trigger_label = zh(locale, "触发时机", "Trigger")
        inputs_label = zh(locale, "需要的输入", "Inputs")
        output_label = zh(locale, "输出结果", "Output")
        question_label = zh(locale, "主动询问用语", "Proactive question")
        no_input_label = zh(locale, "没有新输入时的处理", "When there is no new input")
        has_input_label = zh(locale, "有新输入时的处理", "When new input arrives")
        constraints_label = zh(locale, "约束与原则", "Constraints")
        lines = [
            zh(locale, ZH["review_intro"].format(name=agent["name"]), f"{agent['name']} is ready for review:"),
            "",
            f"1. {mission_label}: {p.get('goal') or '-'}",
            f"2. {user_label}: {p.get('primaryUser') or '-'}",
            f"3. {trigger_label}: {p.get('trigger') or '-'}",
            f"4. {inputs_label}: {p.get('inputs') or '-'}",
            f"5. {output_label}: {p.get('output') or '-'}",
            f"6. {question_label}: {p.get('checkPrompt') or '-'}",
            f"7. {no_input_label}: {p.get('noInputAction') or '-'}",
            f"8. {has_input_label}: {p.get('hasInputAction') or '-'}",
            f"9. {constraints_label}: {p.get('constraints') or '-'}",
            "",
            zh(locale, ZH["review_confirm"], "Reply with 'confirm' or 'OK' to create it, or tell me what to change."),
        ]
        return "\n".join(lines)

    def move_to_review(self, agent: dict, runtime: RuntimeBridge, locale: str) -> str:
        agent["status"] = "review"
        agent["updatedAt"] = self.now_iso()
        agent["artifact"] = self.build_artifact(agent, locale)
        self.append_action(runtime, f"Prepared review summary for '{agent['name']}'.")
        self.save_store(self.get_store(runtime), runtime)
        return self.build_review_summary(agent, locale)

    def update_review_agent(self, agent: dict, message: str, runtime: RuntimeBridge, locale: str) -> str:
        patch = self.extract_update_patch(message, locale, runtime)
        changed = False
        for key, value in patch.items():
            if value:
                agent.setdefault("profile", {})[key] = value
                if key == "name":
                    agent["name"] = value
                changed = True
        agent["updatedAt"] = self.now_iso()
        self.save_store(self.get_store(runtime), runtime)
        self.cortex(runtime).record_event(agent, "review_updated", {"message": message.strip()})
        prefix = zh(locale, ZH["updated_review"], "I updated the draft. Please review it again:")
        return prefix + "\n\n" + self.build_review_summary(agent, locale)

    def should_route_intake_to_builder(self, message: str, locale: str, runtime: RuntimeBridge) -> bool:
        classified = self.classify_agent_message(message, locale, runtime)
        action = str(classified.get("action", "")).strip()
        try:
            confidence = float(classified.get("confidence", 0.0) or 0.0)
        except (TypeError, ValueError):
            confidence = 0.0
        builder_actions = {"create_agent", "continue_collecting", "review_blueprint", "edit_agent", "show_agent", "general_agent_help"}
        if action in builder_actions and confidence >= 0.6:
            return True
        if self.is_create_request(message) or self.is_edit_request(message):
            return True
        if self.get_collecting_agent(runtime) or self.get_review_agent(runtime):
            return True
        return self.matches_activation_phrase(message, locale) or self.looks_like_agent_topic(message)

    def looks_like_negative_report(self, message: str) -> bool:
        lowered = message.lower()
        return any(token in message for token in ["\u6ca1\u6709", "\u65e0\u65b0", "\u6ca1\u6709\u65b0\u8d26\u5355", "\u6ca1\u6709\u65b0\u5185\u5bb9"]) or any(token in lowered for token in ["no bill", "nothing new", "no new"])

    def looks_like_positive_report(self, message: str) -> bool:
        lowered = message.lower()
        if "?" in message or "？" in message:
            return False
        if any(token in message for token in ["告诉我", "请告诉我", "多少", "总额", "总的金额", "累计", "目前为止", "到目前为止", "统计", "汇总"]):
            return False
        if any(token in lowered for token in ["tell me", "how much", "total", "summary", "so far", "until now"]):
            return False
        zh_tokens = ["发票", "收据", "截图", "图片", "上传", "记录这笔", "记到账单", "记一笔", "新增消费"]
        en_tokens = ["invoice", "receipt", "screenshot", "image", "uploaded", "record this", "add expense", "log expense"]
        return any(token in message for token in zh_tokens) or any(token in lowered for token in en_tokens)

    def record_agent_event(self, agent: dict, kind: str, message: str, runtime: RuntimeBridge) -> dict:
        stamp = self.now_iso()
        item = {
            "id": f"{agent['id']}-{kind}-{stamp.replace(':', '').replace('-', '')}",
            "kind": kind,
            "message": message.strip(),
            "createdAt": stamp,
        }
        agent.setdefault("records", []).insert(0, item)
        del agent["records"][20:]
        agent["updatedAt"] = stamp
        summary = f"Recorded {kind} update for '{agent['name']}'."
        self.append_action(runtime, summary)
        self.save_store(self.get_store(runtime), runtime)
        self.cortex(runtime).record_event(agent, kind, {"message": message.strip()})
        return item

    def record_agent_payload(self, agent: dict, kind: str, payload: dict, runtime: RuntimeBridge) -> dict:
        stamp = self.now_iso()
        item = {
            "id": f"{agent['id']}-{kind}-{stamp.replace(':', '').replace('-', '')}",
            "kind": kind,
            "message": str(payload.get("message", "")).strip(),
            "createdAt": stamp,
            "payload": payload,
        }
        agent.setdefault("records", []).insert(0, item)
        del agent["records"][20:]
        agent["updatedAt"] = stamp
        self.append_action(runtime, f"Recorded {kind} payload for '{agent['name']}'.")
        self.save_store(self.get_store(runtime), runtime)
        self.cortex(runtime).record_event(agent, kind, payload)
        return item

    def artifacts_root(self, runtime: RuntimeBridge) -> Path:
        path = runtime.root / "generated" / "custom-agents"
        path.mkdir(parents=True, exist_ok=True)
        return path

    def english_slug(self, value: str, fallback: str = "artifact") -> str:
        text = str(value or "").strip().lower()
        token_map = [
            ("家庭", "family"),
            ("账单", "bills"),
            ("消费", "expenses"),
            ("费用", "costs"),
            ("表格", "table"),
            ("文档", "document"),
            ("文件", "file"),
            ("图片", "image"),
            ("海报", "poster"),
            ("学习", "study"),
            ("计划", "plan"),
            ("提醒", "reminder"),
            ("日程", "schedule"),
            ("旅行", "travel"),
            ("菜单", "menu"),
            ("采购", "shopping"),
            ("预算", "budget"),
            ("报告", "report"),
        ]
        parts = [english for token, english in token_map if token in text]
        ascii_tokens = re.findall(r"[a-z0-9]+", text)
        parts.extend(ascii_tokens)
        deduped = []
        for part in parts:
            if part and part not in deduped:
                deduped.append(part)
        slug = "-".join(deduped).strip("-")
        return slug or fallback

    def safe_artifact_slug(self, value: str) -> str:
        return self.english_slug(value, "artifact")

    def looks_like_artifact_request(self, message: str) -> bool:
        lowered = str(message or "").lower()
        nouns = [
            "表格", "excel", "xlsx", "sheet", "spreadsheet", "工作表", "报表",
            "文件", "文档", "word", "docx", "document", "report", "note",
            "图片", "图像", "海报", "封面", "image", "png", "jpg", "poster",
        ]
        request_markers = [
            "生成", "做", "创建", "制作", "导出", "输出", "整理成", "给我", "帮我", "来个", "做个", "做一",
            "generate", "create", "make", "draft", "export", "please", "can you",
        ]
        return any(token in message or token in lowered for token in nouns) and any(
            token in message or token in lowered for token in request_markers
        )

    def extract_artifact_plan(self, message: str) -> list[str]:
        lowered = str(message or "").lower()
        plan = []
        spreadsheet_tokens = ["表格", "excel", "xlsx", "sheet", "spreadsheet", "工作表", "报表"]
        document_tokens = ["文件", "文档", "word", "docx", "document", "report", "note", "说明"]
        image_tokens = ["图片", "图像", "海报", "封面", "image", "png", "jpg", "poster", "配图"]
        if any(token in message or token in lowered for token in spreadsheet_tokens):
            plan.append("spreadsheet")
        if any(token in message or token in lowered for token in document_tokens):
            plan.append("document")
        if any(token in message or token in lowered for token in image_tokens):
            plan.append("image")
        return plan

    def build_artifact_label(self, agent: dict, message: str) -> str:
        tokens = re.findall(r"[A-Za-z0-9\u4e00-\u9fff]+", str(message or ""))
        topic = "".join(tokens[:6]).strip()
        if not topic:
            topic = str(agent.get("name") or "homehub-artifact").strip()
        return topic[:48]

    def create_spreadsheet_artifact(self, agent: dict, message: str, runtime: RuntimeBridge, stamp: str, label: str) -> dict:
        artifact_name = f"{stamp}-{self.safe_artifact_slug(label)}.xlsx"
        path = self.artifacts_root(runtime) / artifact_name
        try:
            from openpyxl import Workbook

            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Summary"
            sheet.append(["Agent", "Generated At", "Request", "Goal"])
            sheet.append([
                str(agent.get("name") or "Custom Agent"),
                self.now_iso(),
                str(message or "").strip(),
                str(agent.get("profile", {}).get("goal") or "").strip(),
            ])
            sheet.append([])
            sheet.append(["Suggested Columns"])
            sheet.append(["Task"])
            sheet.append(["Owner"])
            sheet.append(["Status"])
            sheet.append(["Due Date"])
            workbook.save(path)
        except Exception:
            fallback = path.with_suffix(".csv")
            agent_name = str(agent.get("name") or "Custom Agent").replace('"', '""')
            request_text = str(message or "").replace('"', '""')
            goal_text = str(agent.get("profile", {}).get("goal") or "").replace('"', '""')
            fallback.write_text(
                "Agent,Generated At,Request,Goal\n"
                f"\"{agent_name}\","
                f"\"{self.now_iso()}\","
                f"\"{request_text}\","
                f"\"{goal_text}\"\n",
                encoding="utf-8",
            )
            path = fallback
        return {
            "kind": "spreadsheet",
            "label": f"{label} table",
            "fileName": path.name,
            "path": str(path.relative_to(runtime.root)),
            "url": f"/generated/custom-agents/{path.name}",
        }

    def create_document_artifact(self, agent: dict, message: str, runtime: RuntimeBridge, stamp: str, label: str) -> dict:
        artifact_name = f"{stamp}-{self.safe_artifact_slug(label)}.docx"
        path = self.artifacts_root(runtime) / artifact_name
        try:
            from docx import Document

            document = Document()
            document.add_heading(label, 0)
            document.add_paragraph(f"Agent: {agent.get('name') or 'Custom Agent'}")
            document.add_paragraph(f"Generated at: {self.now_iso()}")
            document.add_paragraph(f"Request: {str(message or '').strip()}")
            goal = str(agent.get("profile", {}).get("goal") or "").strip()
            if goal:
                document.add_paragraph(f"Goal: {goal}")
            document.add_paragraph("Draft outline")
            document.add_paragraph("Context and purpose", style="List Bullet")
            document.add_paragraph("Main points to cover", style="List Bullet")
            document.add_paragraph("Next actions", style="List Bullet")
            document.save(path)
        except Exception:
            fallback = path.with_suffix(".txt")
            fallback.write_text(
                f"{label}\n\n"
                f"Agent: {agent.get('name') or 'Custom Agent'}\n"
                f"Generated at: {self.now_iso()}\n"
                f"Request: {str(message or '').strip()}\n"
                f"Goal: {str(agent.get('profile', {}).get('goal') or '').strip()}\n",
                encoding="utf-8",
            )
            path = fallback
        return {
            "kind": "document",
            "label": f"{label} document",
            "fileName": path.name,
            "path": str(path.relative_to(runtime.root)),
            "url": f"/generated/custom-agents/{path.name}",
        }

    def create_image_artifact(self, agent: dict, message: str, runtime: RuntimeBridge, stamp: str, label: str) -> dict:
        artifact_name = f"{stamp}-{self.safe_artifact_slug(label)}.png"
        path = self.artifacts_root(runtime) / artifact_name
        try:
            from PIL import Image, ImageDraw, ImageFont

            image = Image.new("RGB", (1280, 720), "#0f1724")
            draw = ImageDraw.Draw(image)
            draw.rounded_rectangle((48, 48, 1232, 672), radius=32, fill="#13263b", outline="#55d0b1", width=3)
            font_title = ImageFont.load_default()
            font_body = ImageFont.load_default()
            draw.text((92, 110), label[:42], fill="#f6fbff", font=font_title)
            draw.text((92, 190), f"Agent: {agent.get('name') or 'Custom Agent'}", fill="#b8f6e5", font=font_body)
            draw.text((92, 240), f"Request: {str(message or '').strip()[:110]}", fill="#d7e7f5", font=font_body)
            goal = str(agent.get("profile", {}).get("goal") or "").strip()
            if goal:
                draw.text((92, 290), f"Goal: {goal[:110]}", fill="#d7e7f5", font=font_body)
            draw.text((92, 580), f"Generated by HomeHub at {self.now_iso()}", fill="#7fb6d9", font=font_body)
            image.save(path)
        except Exception:
            fallback = path.with_suffix(".svg")
            fallback.write_text(
                (
                    "<svg xmlns=\"http://www.w3.org/2000/svg\" width=\"1280\" height=\"720\" viewBox=\"0 0 1280 720\">"
                    "<rect width=\"1280\" height=\"720\" fill=\"#0f1724\"/>"
                    "<rect x=\"48\" y=\"48\" width=\"1184\" height=\"624\" rx=\"32\" fill=\"#13263b\" stroke=\"#55d0b1\" stroke-width=\"3\"/>"
                    f"<text x=\"92\" y=\"120\" fill=\"#f6fbff\" font-size=\"34\" font-family=\"sans-serif\">{escape(label[:42])}</text>"
                    f"<text x=\"92\" y=\"200\" fill=\"#b8f6e5\" font-size=\"20\" font-family=\"sans-serif\">Agent: {escape(str(agent.get('name') or 'Custom Agent'))}</text>"
                    f"<text x=\"92\" y=\"250\" fill=\"#d7e7f5\" font-size=\"18\" font-family=\"sans-serif\">Request: {escape(str(message or '').strip()[:110])}</text>"
                    f"<text x=\"92\" y=\"300\" fill=\"#d7e7f5\" font-size=\"18\" font-family=\"sans-serif\">Goal: {escape(str(agent.get('profile', {}).get('goal') or '').strip()[:110])}</text>"
                    f"<text x=\"92\" y=\"600\" fill=\"#7fb6d9\" font-size=\"16\" font-family=\"sans-serif\">Generated by HomeHub at {escape(self.now_iso())}</text>"
                    "</svg>"
                ),
                encoding="utf-8",
            )
            path = fallback
        return {
            "kind": "image",
            "label": f"{label} image",
            "fileName": path.name,
            "path": str(path.relative_to(runtime.root)),
            "url": f"/generated/custom-agents/{path.name}",
        }

    def handle_artifact_generation(self, agent: dict, message: str, runtime: RuntimeBridge, locale: str) -> dict | None:
        if not self.looks_like_artifact_request(message):
            return None
        if str(agent.get("generatedFeatureId", "")).strip() and any(
            token in message for token in ["消费", "账单", "费用", "支出", "记录", "明细", "统计", "总额"]
        ):
            return None
        plan = self.extract_artifact_plan(message)
        if not plan:
            return None
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        label = self.build_artifact_label(agent, message)
        artifacts = []
        for item in plan:
            if item == "spreadsheet":
                artifacts.append(self.create_spreadsheet_artifact(agent, message, runtime, stamp, label))
            elif item == "document":
                artifacts.append(self.create_document_artifact(agent, message, runtime, stamp, label))
            elif item == "image":
                artifacts.append(self.create_image_artifact(agent, message, runtime, stamp, label))
        payload = {"message": message.strip(), "artifacts": artifacts}
        self.record_agent_payload(agent, "artifact_generation", payload, runtime)
        if locale == "zh-CN":
            labels = "、".join(item["fileName"] for item in artifacts)
            reply = f"我已经为 {agent['name']} 生成好了可下载的产物：{labels}。"
        else:
            labels = ", ".join(item["fileName"] for item in artifacts)
            reply = f"I generated downloadable artifacts for {agent['name']}: {labels}."
        return {"reply": reply, "uiAction": self.studio_ui_action(agent), "artifacts": artifacts}

    def handle_operational_report(self, agent: dict, message: str, runtime: RuntimeBridge, locale: str) -> dict | None:
        generated_feature_id = str(agent.get("generatedFeatureId", "")).strip()
        if generated_feature_id:
            forwarded = runtime.call_feature(generated_feature_id, {"mode": "voice", "message": message}, locale)
            if isinstance(forwarded, dict):
                forwarded_reply = str(forwarded.get("reply") or forwarded.get("message") or "").strip()
                artifacts = forwarded.get("artifacts", []) if isinstance(forwarded.get("artifacts"), list) else []
                if forwarded_reply or artifacts:
                    return {
                        "reply": forwarded_reply or zh(locale, f"{agent['name']} 已经处理好了这次请求。", f"{agent['name']} handled that request."),
                        "uiAction": self.studio_ui_action(agent),
                        "featureResult": forwarded,
                        "artifacts": artifacts,
                    }
        if self.looks_like_negative_report(message):
            self.record_agent_event(agent, "no_input", message, runtime)
            reply = zh(locale, f"\u5df2\u7ecf\u8bb0\u5f55\uff1a{agent['name']}\u672c\u8f6e\u6ca1\u6709\u65b0\u8d26\u5355\uff0c\u6211\u5c31\u8df3\u8fc7\u8fd9\u4e00\u8f6e\u3002", f"I recorded that {agent['name']} has no new input this round.")
            return {"reply": reply, "uiAction": self.studio_ui_action(agent)}
        if self.looks_like_positive_report(message):
            self.record_agent_event(agent, "has_input", message, runtime)
            reply = zh(locale, f"\u5df2\u7ecf\u8bb0\u5f55\u8fd9\u6b21\u53d1\u6765\u7684\u5185\u5bb9\uff0c{agent['name']}\u540e\u7eed\u4f1a\u6309\u84dd\u56fe\u53bb\u5f52\u6863\u3001\u6c47\u603b\u548c\u63d0\u9192\u3002", f"I recorded this new input for {agent['name']} and will handle it according to the blueprint.")
            return {"reply": reply, "uiAction": self.studio_ui_action(agent)}
        return None

    def handle_operational_payload(self, agent: dict, message: str, attachments: list[dict], runtime: RuntimeBridge, locale: str) -> dict:
        attachment = attachments[0] if attachments else {}
        ocr_result = runtime.call_feature(
            "document-ocr",
            {
                "mode": "voice",
                "action": "analyze_attachment",
                "message": message.strip(),
                "attachment": attachment,
                "agent": {
                    "name": agent.get("name", ""),
                    "profile": agent.get("profile", {}),
                    "generatedFeatureId": agent.get("generatedFeatureId", ""),
                },
            },
            locale,
        )
        analysis = {}
        if isinstance(ocr_result, dict) and isinstance(ocr_result.get("analysis"), dict):
            analysis = dict(ocr_result.get("analysis", {}))
        summary = str(analysis.get("summary", "")).strip() or (message.strip() if message.strip() else "Uploaded an image for the agent to review.")
        payload = {
            "message": message.strip(),
            "attachments": [
                {
                    "name": str(item.get("name", "")).strip(),
                    "mimeType": str(item.get("mimeType", "")).strip(),
                    "sizeBytes": int(item.get("sizeBytes", 0) or 0),
                }
                for item in attachments
            ],
            "analysis": analysis,
        }
        self.record_agent_payload(agent, "attachment_input", payload, runtime)
        generated_feature_id = str(agent.get("generatedFeatureId", "")).strip()
        extracted_expenses = self.extract_expenses_from_analysis(analysis, message)
        recommended_action = str(analysis.get("recommendedAction", "")).strip().lower()
        should_record_expense = generated_feature_id and extracted_expenses and (
            recommended_action == "record_expense" or self.is_expense_or_bill_agent(agent, message)
        )
        if should_record_expense:
            forwarded = runtime.call_feature(
                generated_feature_id,
                {
                    "mode": "voice",
                    "action": "import_expenses",
                    "message": message.strip(),
                    "expenses": extracted_expenses,
                    "analysis": analysis,
                },
                locale,
            )
            if isinstance(forwarded, dict) and forwarded.get("ok"):
                forwarded_reply = str(forwarded.get("reply") or "").strip()
                return {
                    "reply": forwarded_reply or zh(locale, f"{agent['name']} 已经把图片中的金额记到账单里了。", f"{agent['name']} recorded the amounts from the image."),
                    "uiAction": self.studio_ui_action(agent),
                    "analysis": analysis,
                    "featureResult": forwarded,
                    "artifacts": forwarded.get("artifacts", []) if isinstance(forwarded.get("artifacts"), list) else [],
                }
        if self.is_expense_or_bill_agent(agent, message) and not extracted_expenses:
            no_amount_reply = zh(
                locale,
                f"我看过这张图片了，但还没有从里面识别到可入账的金额，所以这次没有记到账单里。你可以换一张更清晰的图片，或者直接补一句金额。",
                "I reviewed the image, but I could not detect a bill amount to record yet. Try a clearer image or add the amount in text.",
            )
            return {"reply": no_amount_reply, "uiAction": self.studio_ui_action(agent), "analysis": analysis}
        follow_up = str(analysis.get("followUpQuestion", "")).strip()
        suggested = str(analysis.get("suggestedAction", "")).strip()
        if locale == "zh-CN":
            reply = f"\u6211\u5df2\u7ecf\u628a\u8fd9\u4efd\u4e0a\u4f20\u5185\u5bb9\u8bb0\u5f55\u5230 {agent['name']}\u3002\u6458\u8981\uff1a{summary}"
            if suggested:
                reply += f"\u3002\u5efa\u8bae\u5904\u7406\uff1a{suggested}"
            if follow_up:
                reply += f"\u3002\u4e0b\u4e00\u6b65\u53ef\u4ee5\u8865\u5145\uff1a{follow_up}"
            return {"reply": reply, "uiAction": self.studio_ui_action(agent), "analysis": analysis}
        reply = f"I saved this uploaded image into {agent['name']}. Summary: {summary}"
        if suggested:
            reply += f" Suggested action: {suggested}."
        if follow_up:
            reply += f" Follow-up: {follow_up}"
        return {"reply": reply, "uiAction": self.studio_ui_action(agent), "analysis": analysis}

    def handle_network_lookup(self, agent: dict, query: str, runtime: RuntimeBridge, locale: str) -> dict:
        if not self.blueprint_allows_network(agent):
            reply = zh(locale, f"{agent['name']} \u8fd9\u4efd\u84dd\u56fe\u8fd8\u6ca1\u6709\u5f00\u542f\u8054\u7f51\u67e5\u8be2\u3002\u53ef\u4ee5\u5148\u5728\u84dd\u56fe\u91cc\u8865\u5145\u201c\u5141\u8bb8\u8054\u7f51\u67e5\u8be2\u201d\u3002", f"{agent['name']} is not allowed to use network lookup yet. Update the blueprint to allow it first.")
            return {"reply": reply, "uiAction": self.studio_ui_action(agent), "lookup": {"ok": False, "error": "network_lookup_disabled"}}
        if not runtime.network_lookup:
            reply = zh(locale, "\u5f53\u524d runtime \u8fd8\u6ca1\u6709\u542f\u7528\u8054\u7f51\u67e5\u8be2\u80fd\u529b\u3002", "The current runtime does not have network lookup enabled yet.")
            return {"reply": reply, "uiAction": self.studio_ui_action(agent), "lookup": {"ok": False, "error": "network_lookup_unavailable"}}
        lookup = runtime.network_lookup(
            query,
            locale,
            str(agent.get("profile", {}).get("lookupPolicy", "official-only")).strip() or "official-only",
            self.preferred_sources(agent),
            None,
        )
        self.record_agent_payload(agent, "network_lookup", {"message": query.strip(), "lookup": lookup}, runtime)
        if locale == "zh-CN":
            if not lookup.get("ok"):
                reply = f"\u6211\u8bd5\u8fc7\u53d7\u63a7\u8054\u7f51\u67e5\u8be2\uff0c\u4f46\u8fd9\u6b21\u6ca1\u6709\u62ff\u5230\u53ef\u7528\u7ed3\u679c\uff1a{lookup.get('error', 'unknown error')}\u3002"
            else:
                answer = str(lookup.get("answer", "")).strip() or "\u6211\u627e\u5230\u4e86\u4e00\u4e9b\u5916\u90e8\u4fe1\u606f\u3002"
                reply = f"\u6211\u5df2\u7ecf\u4e3a {agent['name']} \u6267\u884c\u4e86\u53d7\u63a7\u8054\u7f51\u67e5\u8be2\u3002{answer}"
        else:
            if not lookup.get("ok"):
                reply = f"I tried a controlled network lookup, but it did not return usable results: {lookup.get('error', 'unknown error')}."
            else:
                answer = str(lookup.get("answer", "")).strip() or "I found some external information."
                reply = f"I ran a controlled network lookup for {agent['name']}. {answer}"
        return {"reply": reply, "uiAction": self.studio_ui_action(agent), "lookup": lookup}

    def render_generated_feature(self, agent: dict) -> str:
        p = agent.get("profile", {})
        feature_id = self.generated_feature_id(agent)
        feature_name = p.get("name") or agent.get("name") or "Generated Feature"
        summary = p.get("goal") or "Generated from HomeHub blueprint."
        voice_summary = p.get("output") or summary
        api_path = f"/api/{feature_id}"
        api_items_path = f"{api_path}/items"
        api_run_path = f"{api_path}/run"
        blueprint = json.dumps(
            {
                "name": feature_name,
                "goal": p.get("goal", ""),
                "primaryUser": p.get("primaryUser", ""),
                "trigger": p.get("trigger", ""),
                "inputs": p.get("inputs", ""),
                "output": p.get("output", ""),
                "checkPrompt": p.get("checkPrompt", ""),
                "noInputAction": p.get("noInputAction", ""),
                "hasInputAction": p.get("hasInputAction", ""),
                "allowNetworkLookup": p.get("allowNetworkLookup", ""),
                "preferredSources": p.get("preferredSources", ""),
                "lookupPolicy": p.get("lookupPolicy", ""),
                "constraints": p.get("constraints", ""),
            },
            ensure_ascii=True,
            indent=4,
        )
        zh_reply = json.dumps("\u5df2\u7ecf\u6709 feature \u6a21\u677f\u4e86\uff0c\u4f46\u5177\u4f53\u6267\u884c\u903b\u8f91\u8fd8\u9700\u8981\u8865\u5b8c\u3002\u5f53\u524d\u76ee\u6807\u662f\uff1a{goal}\u3002", ensure_ascii=True)
        placeholder_item = json.dumps(
            {
                "id": "replace-me",
                "title": "Replace with a real record structure",
                "summary": "Example storage item created by the scaffold.",
                "createdAt": "2026-01-01T09:00",
            },
            ensure_ascii=True,
            indent=4,
        )
        return f'''from __future__ import annotations

import json
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from ..base import HomeHubFeature, RuntimeBridge


BLUEPRINT = {blueprint}
API_ROOT = {json.dumps(api_path, ensure_ascii=True)}
API_ITEMS = {json.dumps(api_items_path, ensure_ascii=True)}
API_RUN = {json.dumps(api_run_path, ensure_ascii=True)}
STORAGE_TEMPLATE = {placeholder_item}


class Feature(HomeHubFeature):
    feature_id = {json.dumps(feature_id, ensure_ascii=True)}
    feature_name = {json.dumps(feature_name, ensure_ascii=True)}
    version = "1.0.0"

    def descriptor(self) -> dict:
        data = super().descriptor()
        data["summary"] = {json.dumps(summary, ensure_ascii=True)}
        data["blueprint"] = BLUEPRINT
        data["api"] = [API_ROOT, API_ITEMS, API_RUN]
        data["voiceIntents"] = self.voice_intents()
        return data

    def voice_intents(self) -> list[dict]:
        return [{{"id": f"{{self.feature_id}}-intent", "name": self.feature_name, "summary": {json.dumps(voice_summary, ensure_ascii=True)}}}]

    def storage_path(self, runtime: RuntimeBridge) -> Path:
        data_dir = runtime.root / "data"
        data_dir.mkdir(parents=True, exist_ok=True)
        return data_dir / f"{{self.feature_id}}.json"

    def default_store(self) -> dict:
        now = self.now_iso()
        return {{
            "settings": {{"enabled": True}},
            "items": [],
            "recentActions": [
                {{
                    "id": f"{{self.feature_id}}-ready",
                    "summary": f"{{self.feature_name}} scaffold is ready for implementation.",
                    "createdAt": now,
                }}
            ],
            "lastRun": "",
        }}

    def load_store(self, runtime: RuntimeBridge) -> dict:
        path = self.storage_path(runtime)
        if not path.exists():
            store = self.default_store()
            self.save_store(store, runtime)
            return store
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            store = self.default_store()
            self.save_store(store, runtime)
            return store
        if not isinstance(data, dict):
            return self.default_store()
        return {{
            "settings": data.get("settings", {{}}) if isinstance(data.get("settings"), dict) else {{}},
            "items": data.get("items", []) if isinstance(data.get("items"), list) else [],
            "recentActions": data.get("recentActions", []) if isinstance(data.get("recentActions"), list) else [],
            "lastRun": str(data.get("lastRun", "")),
        }}

    def save_store(self, store: dict, runtime: RuntimeBridge) -> None:
        self.storage_path(runtime).write_text(json.dumps(store, ensure_ascii=False, indent=2), encoding="utf-8")

    def on_refresh(self, runtime: RuntimeBridge) -> None:
        runtime.state[self.feature_id] = self.load_store(runtime)

    def reset(self, runtime: RuntimeBridge) -> None:
        store = self.default_store()
        runtime.state[self.feature_id] = store
        self.save_store(store, runtime)

    def get_store(self, runtime: RuntimeBridge) -> dict:
        store = runtime.state.get(self.feature_id)
        if not isinstance(store, dict):
            store = self.load_store(runtime)
            runtime.state[self.feature_id] = store
        return store

    def now_iso(self) -> str:
        return datetime.now().replace(second=0, microsecond=0).isoformat(timespec="minutes")

    def append_action(self, runtime: RuntimeBridge, summary: str) -> None:
        store = self.get_store(runtime)
        stamp = self.now_iso()
        store.setdefault("recentActions", []).insert(0, {{
            "id": f"{{self.feature_id}}-{{stamp.replace(':', '').replace('-', '')}}",
            "summary": summary,
            "createdAt": stamp,
        }})
        del store["recentActions"][12:]
        self.save_store(store, runtime)

    def run_feature(self, runtime: RuntimeBridge, source: str, payload: dict | None = None) -> dict:
        store = self.get_store(runtime)
        stamp = self.now_iso()
        store["lastRun"] = stamp
        summary = f"{{self.feature_name}} was triggered from {{source}}."
        if payload:
            summary += f" Payload keys: {{', '.join(sorted(payload.keys()))}}."
        self.append_action(runtime, summary)
        self.save_store(store, runtime)
        return {{
            "status": "stub",
            "message": "Replace run_feature() with the real workflow.",
            "lastRun": stamp,
            "blueprint": BLUEPRINT,
        }}

    def handle_voice_chat(self, message: str, locale: str, runtime: RuntimeBridge) -> dict | None:
        goal = BLUEPRINT.get("goal", self.feature_name)
        result = self.run_feature(runtime, "voice", {{"message": message}})
        reply = {zh_reply}.format(goal=goal) if locale == "zh-CN" else f"{{self.feature_name}} already has a feature scaffold, but its execution logic still needs to be implemented. Current goal: {{goal}}."
        return {{"reply": reply, "blueprint": BLUEPRINT, "result": result}}

    def enhance_household_modules(self, modules: list[dict], locale: str, runtime: RuntimeBridge) -> list[dict]:
        current = deepcopy(modules)
        store = self.get_store(runtime)
        current.append({{
            "id": self.feature_id,
            "name": self.feature_name,
            "summary": BLUEPRINT.get("goal", ""),
            "state": "active" if store.get("lastRun") else "ready",
            "actionLabel": "Open",
        }})
        return current

    def dashboard_payload(self, locale: str, runtime: RuntimeBridge) -> dict:
        store = self.get_store(runtime)
        return {{
            self.feature_id: {{
                "blueprint": BLUEPRINT,
                "status": "scaffold",
                "storagePath": str(self.storage_path(runtime)),
                "itemCount": len(store.get("items", [])),
                "lastRun": store.get("lastRun", ""),
                "recentActions": store.get("recentActions", [])[:4],
            }}
        }}

    def handle_api(self, method: str, path: str, query: dict, body: dict | None, runtime: RuntimeBridge) -> dict | None:
        store = self.get_store(runtime)
        if method == "GET" and path == API_ROOT:
            return {{
                "status": 200,
                "body": {{
                    "featureId": self.feature_id,
                    "featureName": self.feature_name,
                    "blueprint": BLUEPRINT,
                    "status": "scaffold",
                    "storagePath": str(self.storage_path(runtime)),
                    "suggestedRoutes": [API_ROOT, API_ITEMS, API_RUN],
                    "store": store,
                    "implementationNotes": [
                        "Use items[] for persisted domain records.",
                        "Expand the settings map with runtime switches or owner preferences.",
                        "Replace run_feature() with the real orchestration pipeline.",
                    ],
                }},
            }}
        if method == "GET" and path == API_ITEMS:
            return {{"status": 200, "body": {{"items": store.get("items", []), "recentActions": store.get("recentActions", [])[:10]}}}}
        if method == "POST" and path == API_ITEMS:
            payload = body or {{}}
            title = str(payload.get("title", "")).strip() or STORAGE_TEMPLATE["title"]
            item = {{
                "id": f"{{self.feature_id}}-item-{{self.now_iso().replace(':', '').replace('-', '')}}",
                "title": title,
                "summary": str(payload.get("summary", "")).strip() or STORAGE_TEMPLATE["summary"],
                "createdAt": self.now_iso(),
            }}
            store.setdefault("items", []).insert(0, item)
            self.append_action(runtime, f"Captured a scaffold item for {{self.feature_name}}.")
            self.save_store(store, runtime)
            return {{"status": 201, "body": {{"ok": True, "item": item, "items": store.get("items", [])}}}}
        if method == "POST" and path == API_RUN:
            result = self.run_feature(runtime, "api", body or {{}})
            return {{"status": 200, "body": result}}
        return None


def load_feature() -> HomeHubFeature:
    return Feature()
'''

    def sanitize_identifier(self, value: str, fallback: str = "custom_feature") -> str:
        value = self.english_slug(value, fallback).replace("-", "_")
        if not value:
            value = fallback
        return f"feature_{value}" if value[0].isdigit() else value

    def generated_feature_id(self, agent: dict) -> str:
        source = str(agent.get("profile", {}).get("name", "")).strip() or agent.get("id", "custom_feature")
        slug = self.sanitize_identifier(source)
        return slug[:-8] if slug.endswith("_feature") else slug

    def generated_feature_path(self, agent: dict, runtime: RuntimeBridge) -> Path:
        return runtime.root / "features" / "customize" / f"{self.generated_feature_id(agent)}_feature.py"

    def generate_feature_template(self, agent: dict, runtime: RuntimeBridge, overwrite: bool = False) -> dict:
        if agent.get("status") != "complete":
            return {"ok": False, "error": "not_ready", "path": ""}
        path = self.generated_feature_path(agent, runtime)
        path.parent.mkdir(parents=True, exist_ok=True)
        if path.exists() and not overwrite:
            return {"ok": False, "error": "exists", "path": str(path)}
        path.write_text(self.render_generated_feature(agent), encoding="utf-8")
        agent["generatedFeaturePath"] = str(path)
        agent["generatedFeatureId"] = self.generated_feature_id(agent)
        agent["updatedAt"] = self.now_iso()
        self.append_action(runtime, f"Generated feature scaffold for '{agent['name']}' at {path.name}.")
        self.save_store(self.get_store(runtime), runtime)
        self.cortex(runtime).record_event(agent, "feature_generated", {"message": str(path)})
        return {"ok": True, "path": str(path)}

    def studio_ui_action(self, agent: dict | None = None, focus_input: bool = False) -> dict:
        if focus_input:
            return {
                "type": "focus_element",
                "tab": "test",
                "targetId": "test-input",
                "selector": "#test-input",
                "source": "custom-agent-builder",
            }
        if agent:
            return {
                "type": "select_agent",
                "tab": "test",
                "agentId": agent.get("id", ""),
                "selector": f'[data-agent-id="{agent.get("id", "")}"]',
                "source": "custom-agent-builder",
            }
        return {
            "type": "focus_element",
            "tab": "test",
            "targetId": "test-blueprints-title",
            "selector": "#studio-blueprints .remote-target",
            "source": "custom-agent-builder",
        }

    def answer_collecting_agent(self, agent: dict, message: str, runtime: RuntimeBridge, locale: str) -> str:
        if self.is_cancel_request(message):
            agent["status"] = "cancelled"
            agent["updatedAt"] = self.now_iso()
            self.append_action(runtime, f"Cancelled custom agent draft '{agent['name']}'.")
            self.save_store(self.get_store(runtime), runtime)
            self.cortex(runtime).record_event(agent, "draft_cancelled", {"message": message.strip()})
            return zh(locale, ZH["cancel"].format(name=agent["name"]), f"Okay, I stopped working on {agent['name']}.")
        next_item = self.next_question(agent, locale)
        if not next_item:
            return self.move_to_review(agent, runtime, locale)
        key, _question = next_item
        answer = message.strip() or (self.default_name(agent.get("profile", {}).get("goal", ""), locale) if key == "name" else "")
        agent.setdefault("profile", {})[key] = answer
        if key == "name":
            agent["name"] = answer
        agent.setdefault("qaHistory", []).append({"questionKey": key, "answer": answer, "createdAt": self.now_iso()})
        agent["updatedAt"] = self.now_iso()
        follow_up = self.next_question(agent, locale)
        self.save_store(self.get_store(runtime), runtime)
        self.cortex(runtime).record_event(agent, "question_answered", {"message": answer, "questionKey": key})
        if follow_up:
            return follow_up[1]
        return self.move_to_review(agent, runtime, locale)

    def agent_detail_message(self, agent: dict, locale: str) -> str:
        if not agent:
            return zh(locale, ZH["empty"], "No custom agents yet.")
        if agent.get("status") == "review":
            return self.build_review_summary(agent, locale)
        if agent.get("status") == "complete":
            p = agent.get("profile", {})
            reply = zh(locale, ZH["detail"].format(name=agent["name"], goal=p.get("goal") or "\u672a\u586b\u5199", trigger=p.get("trigger") or "\u672a\u586b\u5199", output=p.get("output") or "\u672a\u586b\u5199"), f"{agent['name']} is complete.")
            if locale == "zh-CN" and agent.get("generatedFeaturePath"):
                reply += " " + ZH["generated"].format(name=agent["name"], path=agent["generatedFeaturePath"])
            return reply
        next_item = self.next_question(agent, locale)
        return zh(locale, ZH["collecting"].format(name=agent["name"], question=next_item[1] if next_item else "\u7ee7\u7eed\u8865\u5145\u8d44\u6599"), f"{agent['name']} is still collecting information.")

    def is_general_time_or_weather_query(self, message: str) -> bool:
        lowered = str(message or "").lower()
        weather_tokens = ["天气", "weather"]
        clock_tokens = ["几点", "现在几点", "当前时间", "what time", "current time"]
        return any(token in message or token in lowered for token in weather_tokens + clock_tokens)

    def match_voice_intent(self, message: str, locale: str, runtime: RuntimeBridge) -> dict | None:
        lowered = message.lower()
        collecting = self.get_collecting_agent(runtime)
        review = self.get_review_agent(runtime)
        operational = self.find_matching_operational_agent(message, runtime)
        activation = self.matches_activation_phrase(message, locale)
        classified = self.classify_agent_message(message, locale, runtime)
        classified_action = str(classified.get("action", "")).strip()
        classified_confidence = float(classified.get("confidence", 0.0) or 0.0)
        if classified_action in {
            "create_agent",
            "continue_collecting",
            "review_blueprint",
            "edit_agent",
            "operational_agent",
            "show_agent",
            "general_agent_help",
        } and classified_confidence >= 0.6:
            return {
                "intent": "custom-agent-builder",
                "action": classified_action,
                "score": min(0.99, max(0.72, classified_confidence)),
                "reason": str(classified.get("reason", "")).strip(),
                "targetAgentName": str(classified.get("targetAgentName", "")).strip(),
            }
        if self.should_generate_feature(message):
            return {"intent": "custom-agent-builder", "action": "generate_feature", "score": 0.98}
        if self.is_edit_request(message):
            return {"intent": "custom-agent-builder", "action": "edit_agent", "score": 0.98}
        if self.is_create_request(message):
            return {"intent": "custom-agent-builder", "action": "create_agent", "score": 0.99}
        if activation:
            if self.is_list_request(message):
                return {"intent": "custom-agent-builder", "action": "show_agent", "score": 0.93}
            return {"intent": "custom-agent-builder", "action": "general_agent_help", "score": 0.9}
        if collecting and not self.is_general_time_or_weather_query(message):
            return {"intent": "custom-agent-builder", "action": "continue_collecting", "score": 0.95}
        if review and not self.is_general_time_or_weather_query(message):
            return {"intent": "custom-agent-builder", "action": "review_blueprint", "score": 0.96}
        if operational and not self.looks_like_agent_topic(message) and not self.is_general_time_or_weather_query(message):
            return {"intent": "custom-agent-builder", "action": "operational_agent", "score": 0.91}
        if self.looks_like_agent_topic(message) and self.is_list_request(message):
            return {"intent": "custom-agent-builder", "action": "show_agent", "score": 0.9}
        if self.looks_like_agent_topic(message):
            return {"intent": "custom-agent-builder", "action": "general_agent_help", "score": 0.78}
        return None

    def handle_voice_chat(self, message: str, locale: str, runtime: RuntimeBridge) -> dict | None:
        collecting = self.get_collecting_agent(runtime)
        review = self.get_review_agent(runtime)
        classified = self.classify_agent_message(message, locale, runtime)
        classified_action = str(classified.get("action", "")).strip()
        classified_confidence = float(classified.get("confidence", 0.0) or 0.0)
        classified_target = self.find_agent_by_hint(str(classified.get("targetAgentName", "")).strip(), runtime) if str(classified.get("targetAgentName", "")).strip() else None
        if self.should_generate_feature(message):
            agent = self.find_agent_by_hint(message, runtime) or self.latest_completed_agent(runtime)
            if not agent:
                return {"reply": zh(locale, ZH["missing_feature"], "I could not find a completed blueprint to turn into a feature scaffold yet."), "uiAction": self.studio_ui_action()}
            result = self.generate_feature_template(agent, runtime, overwrite=self.should_overwrite_feature(message))
            if not result["ok"]:
                if result["error"] == "not_ready":
                    return {"reply": zh(locale, ZH["not_ready"].format(name=agent["name"]), f"{agent['name']} is not ready to generate a feature scaffold yet."), "uiAction": self.studio_ui_action(agent)}
                return {"reply": zh(locale, ZH["exists"].format(name=agent["name"], path=result["path"]), f"The feature scaffold already exists at {result['path']}."), "uiAction": self.studio_ui_action(agent)}
            return {"reply": zh(locale, ZH["generated"].format(name=agent["name"], path=result["path"]), f"I generated the feature scaffold for {agent['name']} at {result['path']}."), "uiAction": self.studio_ui_action(agent)}
        if self.is_edit_request(message) or (classified_action == "edit_agent" and classified_confidence >= 0.6):
            target = classified_target or self.find_agent_by_hint(message, runtime) or self.latest_completed_agent(runtime)
            if not target:
                return {"reply": zh(locale, ZH["empty"], "No custom agents yet."), "uiAction": self.studio_ui_action()}
            agent = self.start_editing_agent(target, message, runtime, locale)
            intro = zh(locale, ZH["edit_started"].format(name=target["name"]), f"I started an editable draft based on {target['name']}.")
            return {"reply": intro + "\n\n" + self.build_review_summary(agent, locale), "uiAction": self.studio_ui_action(agent)}
        if self.is_create_request(message):
            draft_profile = self.infer_initial_profile(message, locale, runtime)
            if not self.should_force_create(message):
                similar = self.find_similar_agents(draft_profile, runtime)
                if similar:
                    return {"reply": self.similar_agents_message(draft_profile, similar, locale), "uiAction": self.studio_ui_action(similar[0])}
            agent = self.create_agent(runtime, message, locale)
            q = self.next_question(agent, locale)
            return {"reply": zh(locale, ZH["draft_started"].format(name=agent["name"], question=q[1] if q else ""), f"{agent['name']} started. {q[1] if q else ''}"), "uiAction": self.studio_ui_action(agent)}
        if classified_action == "create_agent" and classified_confidence >= 0.6:
            draft_profile = self.infer_initial_profile(message, locale, runtime)
            similar = self.find_similar_agents(draft_profile, runtime)
            if similar and not self.should_force_create(message):
                return {"reply": self.similar_agents_message(draft_profile, similar, locale), "uiAction": self.studio_ui_action(similar[0])}
            agent = self.create_agent(runtime, message, locale)
            q = self.next_question(agent, locale)
            return {"reply": zh(locale, ZH["draft_started"].format(name=agent["name"], question=q[1] if q else ""), f"{agent['name']} started. {q[1] if q else ''}"), "uiAction": self.studio_ui_action(agent)}
        if collecting and not self.is_general_time_or_weather_query(message):
            reply = self.answer_collecting_agent(collecting, message, runtime, locale)
            target_agent = self.get_collecting_agent(runtime) or self.find_agent_by_hint(collecting.get("name", ""), runtime) or collecting
            return {"reply": reply, "uiAction": self.studio_ui_action(target_agent)}
        if review and not self.is_general_time_or_weather_query(message):
            if self.is_cancel_request(message):
                review["status"] = "cancelled"
                review["updatedAt"] = self.now_iso()
                self.save_store(self.get_store(runtime), runtime)
                return {"reply": zh(locale, ZH["cancel"].format(name=review["name"]), f"Okay, I stopped working on {review['name']}."), "uiAction": self.studio_ui_action(review)}
            if self.is_confirmation_message(message):
                self.complete_agent(review, runtime, locale)
                confirmed = zh(
                    locale,
                    ZH["edit_confirmed"].format(name=review["name"]) if review.get("editMode") else ZH["confirmed"].format(name=review["name"]),
                    f"{review['name']} has been created and confirmed.",
                )
                return {"reply": confirmed, "uiAction": self.studio_ui_action(review)}
            if self.is_revision_message(message) or message.strip():
                return {"reply": self.update_review_agent(review, message, runtime, locale), "uiAction": self.studio_ui_action(review)}
        operational = classified_target or self.find_matching_operational_agent(message, runtime) or self.latest_operational_agent(runtime)
        if operational and not self.looks_like_agent_topic(message) and not self.is_general_time_or_weather_query(message):
            artifact_result = self.handle_artifact_generation(operational, message, runtime, locale)
            if artifact_result:
                return artifact_result
            result = self.handle_operational_report(operational, message, runtime, locale)
            if result:
                return result
        if classified_action == "show_agent" and classified_confidence >= 0.6:
            agent = classified_target or self.find_agent_by_hint(message, runtime) or self.latest_completed_agent(runtime)
            return {"reply": self.agent_detail_message(agent, locale), "uiAction": self.studio_ui_action(agent)} if agent else {"reply": zh(locale, ZH["empty"], "No custom agents yet."), "uiAction": self.studio_ui_action()}
        if classified_action == "general_agent_help" and classified_confidence >= 0.6:
            return {"reply": zh(locale, ZH["help"], "Describe the ongoing job you want this agent to own, and I will ask for whatever information is still missing."), "uiAction": self.studio_ui_action(focus_input=True)}
        if self.matches_activation_phrase(message, locale):
            if self.is_list_request(message):
                agent = self.find_agent_by_hint(message, runtime) or self.latest_completed_agent(runtime)
                return {"reply": self.agent_detail_message(agent, locale), "uiAction": self.studio_ui_action(agent)} if agent else {"reply": zh(locale, ZH["empty"], "No custom agents yet."), "uiAction": self.studio_ui_action()}
            return {"reply": zh(locale, ZH["help"], "Describe the ongoing job you want this agent to own, and I will ask for whatever information is still missing."), "uiAction": self.studio_ui_action(focus_input=True)}
        if self.looks_like_agent_topic(message) and self.is_list_request(message):
            agent = self.find_agent_by_hint(message, runtime)
            return {"reply": self.agent_detail_message(agent, locale), "uiAction": self.studio_ui_action(agent)} if agent else {"reply": zh(locale, ZH["empty"], "No custom agents yet."), "uiAction": self.studio_ui_action()}
        if self.looks_like_agent_topic(message):
            return {"reply": zh(locale, ZH["help"], "Describe the ongoing job you want this agent to own, and I will ask for whatever information is still missing."), "uiAction": self.studio_ui_action(focus_input=True)}
        return None

    def enhance_household_modules(self, modules: list[dict], locale: str, runtime: RuntimeBridge) -> list[dict]:
        current = deepcopy(modules)
        agents = self.sorted_agents(runtime)
        collecting = [item for item in agents if item.get("status") == "collecting"]
        completed = [item for item in agents if item.get("status") == "complete"]
        generated = sum(1 for item in agents if item.get("generatedFeaturePath"))
        if locale == "zh-CN":
            summary = ZH["module_summary_count"].format(total=len(agents), completed=len(completed)) if agents else ZH["module_summary_empty"]
            if generated:
                summary += ZH["module_generated_suffix"].format(count=generated)
            if collecting:
                summary += ZH["module_collecting_suffix"].format(count=len(collecting))
            name = ZH["module_name"]
            action = ZH["module_action_continue"] if collecting else ZH["module_action_create"]
        else:
            summary = f"{len(agents)} custom agents created, {len(completed)} complete."
            name = "Custom Agents"
            action = "Continue Draft" if collecting else "Create by Voice"
        current.append({"id": "custom-agents", "name": name, "summary": summary, "state": "attention" if collecting else ("active" if agents else "ready"), "actionLabel": action})
        return current

    def dashboard_payload(self, locale: str, runtime: RuntimeBridge) -> dict:
        agents = self.sorted_agents(runtime)[:8]
        return {
            "customAgents": agents,
            "customAgentRecentActions": self.get_store(runtime).get("recentActions", [])[:6],
            "customAgentCortex": self.cortex(runtime).summaries_for(agents),
        }

    def handle_api(self, method: str, path: str, query: dict, body: dict | None, runtime: RuntimeBridge) -> dict | None:
        if method == "GET" and path == "/api/custom-agents":
            agents = [self.enrich_agent_runtime_stats(agent, runtime) for agent in self.sorted_agents(runtime)]
            return {
                "status": 200,
                "body": {
                    "items": agents,
                    "recentActions": self.get_store(runtime).get("recentActions", [])[:10],
                    "cortex": self.cortex(runtime).summaries_for(agents),
                },
            }
        if method == "POST" and path == "/api/custom-agents/generate-feature":
            payload = body or {}
            agent = None
            if payload.get("id"):
                for item in self.sorted_agents(runtime):
                    if item.get("id") == str(payload.get("id")):
                        agent = item
                        break
            if agent is None and payload.get("name"):
                agent = self.find_agent_by_hint(str(payload.get("name")), runtime)
            if agent is None:
                agent = self.latest_completed_agent(runtime)
            if agent is None:
                return {"status": 404, "body": {"error": "No completed custom agent blueprint found."}}
            result = self.generate_feature_template(agent, runtime, overwrite=bool(payload.get("overwrite")))
            if not result["ok"]:
                return {"status": 409 if result["error"] == "exists" else 400, "body": {"error": result["error"], "path": result["path"], "agentId": agent.get("id")}}
            return {
                "status": 200,
                "body": {
                    "ok": True,
                    "agentId": agent.get("id"),
                    "featureId": agent.get("generatedFeatureId"),
                    "path": result["path"],
                    "item": agent,
                    "cortex": self.cortex(runtime).get_summary(str(agent.get("id", ""))),
                },
            }
        if method == "GET" and path == "/api/custom-agents/cortex":
            agent_id = str(query.get("id", [""])[0] if isinstance(query.get("id"), list) else query.get("id", "")).strip()
            if not agent_id:
                return {"status": 400, "body": {"error": "id is required"}}
            return {"status": 200, "body": {"item": self.cortex(runtime).get_summary(agent_id)}}
        if method == "POST" and path == "/api/custom-agents/autonomous-plan":
            payload = body or {}
            message = str(payload.get("message", "")).strip()
            locale = str(payload.get("locale", "zh-CN")).strip() or "zh-CN"
            task_type = str(payload.get("taskType", "general_chat")).strip() or "general_chat"
            attachments = payload.get("attachments", []) if isinstance(payload.get("attachments", []), list) else []
            request = self.build_request_context(
                message,
                locale,
                attachments=attachments,
                task_type=task_type,
                requires_network=bool(payload.get("requiresNetwork", False)),
                require_artifacts=bool(payload.get("requireArtifacts", False)),
            )
            plan = self.should_autonomously_create(request, runtime)
            return {"status": 200, "body": {"ok": True, "request": request, "plan": plan}}
        if method == "POST" and path == "/api/custom-agents/intake":
            payload = body or {}
            message = str(payload.get("message", "")).strip()
            locale = str(payload.get("locale", "zh-CN")).strip() or "zh-CN"
            attachments = payload.get("attachments", []) if isinstance(payload.get("attachments", []), list) else []
            if message and not attachments and self.should_route_intake_to_builder(message, locale, runtime):
                result = self.handle_voice_chat(message, locale, runtime) or {}
                if result:
                    return {
                        "status": 200,
                        "body": {
                            "ok": True,
                            "reply": result.get("reply"),
                            "artifacts": result.get("artifacts", []),
                            "uiAction": result.get("uiAction"),
                        },
                    }
            agent = None
            if payload.get("id"):
                for item in self.sorted_agents(runtime):
                    if item.get("id") == str(payload.get("id")):
                        agent = item
                        break
            if agent is None and payload.get("name"):
                agent = self.find_agent_by_hint(str(payload.get("name")), runtime)
            if agent is None:
                agent = self.latest_operational_agent(runtime) or self.latest_completed_agent(runtime)
            if agent is None:
                return {"status": 404, "body": {"error": "No active custom agent found for intake."}}
            if attachments:
                result = self.handle_operational_payload(agent, message, attachments, runtime, locale)
                return {
                    "status": 200,
                    "body": {
                        "ok": True,
                        "agentId": agent.get("id"),
                        "featureId": agent.get("generatedFeatureId"),
                        "reply": result.get("reply"),
                        "analysis": result.get("analysis", {}),
                        "featureResult": result.get("featureResult", {}) if isinstance(result.get("featureResult"), dict) else {},
                        "artifacts": result.get("artifacts", []),
                        "item": agent,
                        "cortex": self.cortex(runtime).get_summary(str(agent.get("id", ""))),
                    },
                }
            artifact_result = self.handle_artifact_generation(agent, message, runtime, locale)
            if artifact_result:
                return {
                    "status": 200,
                    "body": {
                        "ok": True,
                        "agentId": agent.get("id"),
                        "reply": artifact_result.get("reply"),
                        "artifacts": artifact_result.get("artifacts", []),
                        "item": agent,
                        "cortex": self.cortex(runtime).get_summary(str(agent.get("id", ""))),
                    },
                }
            result = self.handle_operational_report(agent, message, runtime, locale)
            if result:
                return {
                    "status": 200,
                    "body": {
                        "ok": True,
                        "agentId": agent.get("id"),
                        "reply": result.get("reply"),
                        "artifacts": result.get("artifacts", []),
                        "item": agent,
                        "cortex": self.cortex(runtime).get_summary(str(agent.get("id", ""))),
                    },
                }
            return {"status": 400, "body": {"error": "No usable intake payload was detected."}}
        if method == "POST" and path == "/api/custom-agents/lookup":
            payload = body or {}
            agent = None
            if payload.get("id"):
                for item in self.sorted_agents(runtime):
                    if item.get("id") == str(payload.get("id")):
                        agent = item
                        break
            if agent is None and payload.get("name"):
                agent = self.find_agent_by_hint(str(payload.get("name")), runtime)
            if agent is None:
                agent = self.latest_completed_agent(runtime)
            if agent is None:
                return {"status": 404, "body": {"error": "No completed custom agent found for network lookup."}}
            query_text = str(payload.get("query", "")).strip()
            locale = str(payload.get("locale", "zh-CN")).strip() or "zh-CN"
            if not query_text:
                return {"status": 400, "body": {"error": "query is required"}}
            result = self.handle_network_lookup(agent, query_text, runtime, locale)
            return {
                "status": 200,
                "body": {
                    "ok": result.get("lookup", {}).get("ok", False),
                    "agentId": agent.get("id"),
                    "reply": result.get("reply"),
                    "lookup": result.get("lookup", {}),
                    "item": agent,
                    "cortex": self.cortex(runtime).get_summary(str(agent.get("id", ""))),
                },
            }
        return None


def load_feature() -> HomeHubFeature:
    return Feature()
